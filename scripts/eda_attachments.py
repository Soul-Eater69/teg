"""Phase 1 EDA — historical IDMT tickets: attachments + ticket usefulness, in one pass.

For each ticket (one Jira fetch) it captures: attachments (metadata + extracted text), the Jira
description, and the linked themes (links of ANY type carrying a Business Value Stream value) and
the Value Streams they map to. From that it profiles both:
  - attachment sizing for condense (counts, file types, sizes, extracted chars, token budget), and
  - ticket usefulness (a ticket is useful when it has theme/VS ground truth): the usefulness split,
    themes-per-ticket distribution, themes among tickets with no attachments, and VS coverage.

Reuses the production attachment extractor + condense selection rules so the numbers match what
ingestion actually sees. The Business Value Stream field id is discovered by name (override with
--field-id).

Usage (Jira + IDP creds; install: uv sync --extra eda --extra extract):
  uv run python scripts/eda_attachments.py tickets.txt --out out/eda/attachments --docx
  # cached to <out>/attachments_raw.json; re-run resumes; analysis-only with --use-cache.

The notebook notebooks/attachments_eda.ipynb imports these functions for an interactive view.
"""

from __future__ import annotations

import argparse
import asyncio
import json
from pathlib import Path

import httpx

from teg.condense.attachment_ranker import is_idea_card, is_supported, select_attachments
from teg.condense.config import CondenseConfig
from teg.config.settings import load_settings
from teg.ingestion.extraction.value_stream_field import parse_value_stream
from teg.integrations.files import build_attachment_extractor

TOKEN_BUDGET = 20_000  # context-headroom flag for the EDA (gpt-5-mini has far more)
DOC_CHAR_BUDGET = CondenseConfig().doc_char_budget  # the real fallback cap (chars), kept in sync


# ---------------------------------------------------------------- collection

def load_ticket_ids(path: str) -> list[str]:
    lines = Path(path).read_text(encoding="utf-8").splitlines()
    return [s.strip() for s in lines if s.strip() and not s.strip().startswith("#")]


def _ext(filename: str) -> str:
    name = filename.lower().strip()
    return name.rsplit(".", 1)[-1] if "." in name else "(none)"


def estimate_tokens(text: str) -> int:
    """Token count via tiktoken when available, else a ~4-chars-per-token estimate."""
    if not text:
        return 0
    try:
        import tiktoken

        return len(tiktoken.get_encoding("cl100k_base").encode(text))
    except Exception:
        return max(1, len(text) // 4)


async def discover_field_id(http: httpx.AsyncClient, api: str, field_name: str) -> str | None:
    resp = await http.get(f"/rest/api/{api}/field")
    resp.raise_for_status()
    wanted = field_name.strip().lower()
    for f in resp.json() or []:
        if str(f.get("name") or "").strip().lower() == wanted:
            return str(f.get("id") or "")
    return None


def _linked_keys(fields: dict) -> list[str]:
    keys: list[str] = []
    for link in fields.get("issuelinks") or []:
        issue = link.get("inwardIssue") or link.get("outwardIssue")
        if isinstance(issue, dict):
            key = str(issue.get("key") or "")
            if key and key not in keys:
                keys.append(key)
    return keys


async def collect(
    ticket_ids: list[str],
    *,
    bvs_field: str,
    concurrency: int = 4,
    ticket_timeout: float = 90.0,
    done_records: list[dict] | None = None,
    checkpoint_path: Path | None = None,
    checkpoint_every: int = 10,
) -> list[dict]:
    """One 'description' record per ticket (carrying ticket-level theme/VS info) plus one record
    per attachment. One Jira fetch per ticket gets summary/description/attachments/issuelinks;
    each linked issue is then read for its Business Value Stream value to count themes.

    Robust to stalls: every fetch/download/extract is bounded by ``ticket_timeout``. Resumable:
    pass ``done_records`` to skip already-collected tickets and ``checkpoint_path`` to flush every
    ``checkpoint_every`` tickets, so a stop loses at most that many and re-running continues.
    """
    settings = load_settings()
    http = httpx.AsyncClient(
        base_url=settings.jira_base_url,
        headers={"Authorization": f"Bearer {settings.jira_token}"},
        timeout=settings.jira_timeout_seconds, verify=settings.jira_verify_ssl,
    )
    api = settings.jira_api_version
    extractor = build_attachment_extractor()
    sem = asyncio.Semaphore(concurrency)
    # A ticket is "done" only if its cached record carries the theme data (new schema). Records
    # from the older attachments-only collector lack themeCount - re-collect those, don't trust them.
    complete = {r["ticketId"] for r in (done_records or [])
                if r.get("kind") == "description" and "themeCount" in r}
    records: list[dict] = [r for r in (done_records or []) if r["ticketId"] in complete]
    todo = [t for t in ticket_ids if t not in complete]
    stale = len(done_records or []) - len(records)
    if stale:
        print(f"  ignoring {stale} cached records without theme data (old schema) - re-collecting")
    progress = {"n": 0, "total": len(todo)}

    async def _issue(key: str, fields: str) -> dict:
        resp = await http.get(f"/rest/api/{api}/issue/{key}", params={"fields": fields})
        resp.raise_for_status()
        return resp.json() or {}

    async def _theme_vs(key: str):
        try:
            li = await asyncio.wait_for(_issue(key, f"summary,{bvs_field}"), timeout=ticket_timeout)
        except Exception:
            return None
        return parse_value_stream((li.get("fields") or {}).get(bvs_field))

    async def _one(ticket_id: str) -> list[dict]:
        try:
            issue = await asyncio.wait_for(
                _issue(ticket_id, "summary,description,attachment,issuelinks"), timeout=ticket_timeout
            )
        except Exception as exc:  # a bad/slow ticket must not abort the batch - record it
            reason = f"{type(exc).__name__}: {exc}" or type(exc).__name__
            print(f"  {ticket_id}: FETCH ERROR {reason}")
            return [{
                "ticketId": ticket_id, "kind": "fetch_error", "filename": "(ticket fetch)",
                "ext": "(error)", "mimeType": "", "sizeBytes": 0, "charCount": 0, "tokenEst": 0,
                "supported": False, "ideaCard": False, "selected": False, "extractError": reason,
                "linkCount": 0, "themeCount": 0, "valueStreamIds": [],
            }]
        f = issue.get("fields") or {}
        atts = [{"filename": str(a.get("filename") or ""), "content": str(a.get("content") or ""),
                 "mime": str(a.get("mimeType") or ""), "size": int(a.get("size") or 0)}
                for a in (f.get("attachment") or [])]
        linked_keys = _linked_keys(f)

        # Linked themes -> Value Streams (any link type; a theme = a link carrying a VS).
        vs_parsed = await asyncio.gather(*(_theme_vs(k) for k in linked_keys)) if linked_keys else []
        vs_ids = [p[1] for p in vs_parsed if p]
        vs_names = [p[0] for p in vs_parsed if p]

        desc = str(f.get("description") or "")
        out: list[dict] = [{
            "ticketId": ticket_id, "kind": "description", "filename": "(description)",
            "ext": "(description)", "mimeType": "", "sizeBytes": len(desc.encode("utf-8")),
            "charCount": len(desc), "tokenEst": estimate_tokens(desc),
            "supported": True, "ideaCard": False, "selected": False, "extractError": "",
            # ticket-level usefulness info lives on the description row (one per ticket)
            "linkCount": len(linked_keys), "themeCount": len(vs_ids),
            "valueStreamIds": vs_ids, "valueStreamNames": vs_names,
        }]
        # Which attachments condense would actually use: idea-card-first, else top-4 supported.
        names = [a["filename"] for a in atts]
        idea = next((n for n in names if is_idea_card(n)), None)
        if idea is not None:
            selected_names = {idea}
        else:
            supported = [n for n in names if is_supported(n)]
            selected_names = set(sorted(supported)[:4]) if supported else set()
        for a in atts:
            fn = a["filename"]
            rec = {
                "ticketId": ticket_id, "kind": "attachment", "filename": fn, "ext": _ext(fn),
                "mimeType": a["mime"], "sizeBytes": a["size"], "charCount": 0, "tokenEst": 0,
                "supported": is_supported(fn), "ideaCard": is_idea_card(fn),
                "selected": fn in selected_names, "extractError": "",
            }
            if rec["supported"] and a["content"]:
                try:
                    dl = await asyncio.wait_for(http.get(a["content"]), timeout=ticket_timeout)
                    dl.raise_for_status()
                    text = await asyncio.wait_for(
                        asyncio.to_thread(extractor.extract, fn, dl.content), timeout=ticket_timeout
                    )
                    rec["charCount"] = len(text or "")
                    rec["tokenEst"] = estimate_tokens(text or "")
                except Exception as exc:
                    rec["extractError"] = f"{type(exc).__name__}: {exc}" or type(exc).__name__
                    print(f"  {ticket_id} / {fn}: EXTRACT ERROR {rec['extractError']}")
            out.append(rec)
        return out

    async def _guarded(ticket_id: str) -> None:
        async with sem:
            out = await _one(ticket_id)
        records.extend(out)
        progress["n"] += 1
        head = next((r for r in out if r["kind"] in ("description", "fetch_error")), {})
        n_att = sum(1 for r in out if r["kind"] == "attachment")
        print(f"[{progress['n']}/{progress['total']}] {ticket_id}: "
              f"{n_att} attach, {head.get('themeCount', 0)} themes")
        if checkpoint_path is not None and progress["n"] % checkpoint_every == 0:
            checkpoint_path.write_text(json.dumps(records, ensure_ascii=False, indent=2), encoding="utf-8")

    try:
        if complete:
            print(f"resuming: {len(complete)} tickets already collected, {len(todo)} to go")
        await asyncio.gather(*(_guarded(t) for t in todo))
    finally:
        await http.aclose()
    if checkpoint_path is not None:  # final flush
        checkpoint_path.write_text(json.dumps(records, ensure_ascii=False, indent=2), encoding="utf-8")
    return records


# ---------------------------------------------------------------- analysis (pandas)

def to_frames(records: list[dict]):
    """Return (attachments_df, tickets_df). Requires pandas (eda extra)."""
    import pandas as pd

    df = pd.DataFrame(records)
    att = df[df["kind"] == "attachment"].copy()
    desc = df[df["kind"] == "description"].copy()

    att["selectedTokens"] = att["tokenEst"] * att.get("selected", False).astype(int)
    per_ticket = att.groupby("ticketId").agg(
        attachmentCount=("filename", "count"),
        supportedCount=("supported", "sum"),
        attachmentBytes=("sizeBytes", "sum"),
        attachmentTokens=("tokenEst", "sum"),
        selectedTokens=("selectedTokens", "sum"),  # only the attachments condense would feed
        attachmentChars=("charCount", "sum"),
        hasIdeaCard=("ideaCard", "max"),
    )
    desc_idx = desc.set_index("ticketId")
    per_ticket["descriptionChars"] = desc_idx["charCount"].reindex(per_ticket.index).fillna(0).astype(int)
    per_ticket["descriptionTokens"] = desc_idx["tokenEst"].reindex(per_ticket.index).fillna(0).astype(int)
    # All attachments (over-count) vs what condense actually ingests (description + selected only).
    per_ticket["combinedTokens"] = per_ticket["attachmentTokens"] + per_ticket["descriptionTokens"]
    per_ticket["condenseInputTokens"] = per_ticket["selectedTokens"] + per_ticket["descriptionTokens"]
    # tickets with attachments listed in the issue but none fetched still appear via desc only:
    only_desc = desc_idx.index.difference(per_ticket.index)
    if len(only_desc):
        import pandas as pd  # noqa: F811
        extra = pd.DataFrame(index=only_desc)
        extra["attachmentCount"] = 0
        extra["supportedCount"] = 0
        extra["attachmentBytes"] = 0
        extra["attachmentTokens"] = 0
        extra["selectedTokens"] = 0
        extra["attachmentChars"] = 0
        extra["hasIdeaCard"] = False
        extra["descriptionChars"] = desc_idx["charCount"].reindex(only_desc).fillna(0).astype(int)
        extra["descriptionTokens"] = desc_idx["tokenEst"].reindex(only_desc).fillna(0).astype(int)
        extra["combinedTokens"] = extra["descriptionTokens"]
        extra["condenseInputTokens"] = extra["descriptionTokens"]
        per_ticket = pd.concat([per_ticket, extra])
    # ticket-level theme/usefulness info from the description rows (one per ticket).
    if "themeCount" in desc.columns:
        per_ticket["themeCount"] = desc_idx["themeCount"].reindex(per_ticket.index).fillna(0).astype(int)
        per_ticket["valueStreamIds"] = desc_idx["valueStreamIds"].reindex(per_ticket.index)
        per_ticket["valueStreamIds"] = per_ticket["valueStreamIds"].apply(lambda v: v if isinstance(v, list) else [])
    else:
        per_ticket["themeCount"] = 0
        per_ticket["valueStreamIds"] = [[] for _ in range(len(per_ticket))]
    return att, per_ticket.reset_index().rename(columns={"index": "ticketId"})


def budget_applied_input(records: list[dict], *, budget: int = DOC_CHAR_BUDGET) -> dict[str, dict]:
    """The TRUE condense input per ticket, applying the real budget (not the full extracted text).

    Mirrors condense: idea-card path = description + idea card in full (uncapped); fallback path =
    description (full) + each selected top-N doc capped at budget/N chars. Returns
    {ticketId: {chars, tokens, path}}; tokens scale each doc's tiktoken estimate to its capped chars.
    """
    by_ticket: dict[str, list[dict]] = {}
    desc: dict[str, dict] = {}
    for r in records:
        if r.get("kind") == "attachment":
            by_ticket.setdefault(r["ticketId"], []).append(r)
        elif r.get("kind") == "description":
            desc[r["ticketId"]] = r

    out: dict[str, dict] = {}
    for tid, d in desc.items():
        atts = by_ticket.get(tid, [])
        selected = [a for a in atts if a.get("selected")]
        idea = next((a for a in selected if a.get("ideaCard")), None)
        chars = int(d.get("charCount", 0))
        tokens = int(d.get("tokenEst", 0))
        if idea is not None:  # idea-card path: full, uncapped
            chars += int(idea.get("charCount", 0))
            tokens += int(idea.get("tokenEst", 0))
            path = "idea_card"
        else:  # fallback: each selected doc capped at budget/N chars
            n = len(selected)
            per_doc = (budget // n) if n else 0
            for a in selected:
                c = int(a.get("charCount", 0))
                capped = min(c, per_doc)
                chars += capped
                tokens += round(int(a.get("tokenEst", 0)) * (capped / c)) if c else 0
            path = "fallback" if n else "description_only"
        out[tid] = {"chars": chars, "tokens": tokens, "path": path}
    return out


def _save(fig, charts_dir: Path, name: str) -> str:
    charts_dir.mkdir(parents=True, exist_ok=True)
    path = charts_dir / f"{name}.png"
    fig.savefig(path, dpi=150, bbox_inches="tight")
    return str(path)


def build_charts(att, tickets, charts_dir: Path, *, records: list[dict] | None = None) -> tuple[dict, dict]:
    """Compute the metric tables + render the charts. Returns (stats, chart_paths)."""
    import matplotlib
    import pandas as pd

    matplotlib.use("Agg")
    import matplotlib.pyplot as plt

    plt.rcParams.update({"figure.figsize": (8, 4.5), "axes.grid": True, "grid.alpha": 0.3})
    stats: dict = {}
    charts: dict = {}

    # 0. attachment presence: tickets WITH vs WITHOUT any attachment
    n = int(len(tickets))
    without = int((tickets["attachmentCount"] == 0).sum())
    with_att = n - without
    stats["attachment_presence"] = {
        "tickets_total": n,
        "tickets_with_attachments": with_att,
        "tickets_without_attachments": without,
        "pct_without": round(100.0 * without / max(1, n), 1),
    }
    fig, ax = plt.subplots(figsize=(5, 4))
    ax.bar(["with attachments", "no attachments"], [with_att, without], color=["#4C78A8", "#E45756"])
    for i, v in enumerate([with_att, without]):
        ax.text(i, v, str(v), ha="center", va="bottom")
    ax.set(title="Tickets with vs without attachments", ylabel="# tickets")
    charts["attachment_presence"] = _save(fig, charts_dir, "attachment_presence"); plt.close(fig)

    # 1. attachments per ticket
    counts = tickets["attachmentCount"]
    stats["attachments_per_ticket"] = {
        "tickets": int(len(tickets)), "mean": round(float(counts.mean()), 2),
        "median": float(counts.median()), "max": int(counts.max()),
        "with_zero": int((counts == 0).sum()),
    }
    fig, ax = plt.subplots()
    counts.value_counts().sort_index().plot(kind="bar", ax=ax, color="#4C78A8")
    ax.set(title="Attachments per ticket", xlabel="# attachments", ylabel="# tickets")
    charts["attachments_per_ticket"] = _save(fig, charts_dir, "attachments_per_ticket"); plt.close(fig)

    # 2. file types
    types = att["ext"].value_counts()
    stats["file_types"] = {k: int(v) for k, v in types.items()}
    stats["most_common_type"] = types.index[0] if len(types) else None
    fig, ax = plt.subplots()
    types.plot(kind="bar", ax=ax, color="#F58518")
    ax.set(title="Attachment file types", xlabel="extension", ylabel="# attachments")
    charts["file_types"] = _save(fig, charts_dir, "file_types"); plt.close(fig)

    # 3. attachment sizes (KB) — per attachment AND total per ticket
    kb = att["sizeBytes"] / 1024.0
    per_ticket_kb = tickets["attachmentBytes"] / 1024.0
    stats["attachment_size_kb"] = {
        # per single attachment
        "per_attachment_mean": round(float(kb.mean()), 1),
        "per_attachment_median": round(float(kb.median()), 1),
        "per_attachment_p90": round(float(kb.quantile(0.9)), 1),
        "per_attachment_max": round(float(kb.max()), 1),
        # total across a ticket's attachments
        "per_ticket_total_mean": round(float(per_ticket_kb.mean()), 1),
        "per_ticket_total_median": round(float(per_ticket_kb.median()), 1),
        "per_ticket_total_max": round(float(per_ticket_kb.max()), 1),
        "all_attachments_total_mb": round(float(att["sizeBytes"].sum()) / 1_048_576, 1),
    }
    fig, ax = plt.subplots()
    ax.hist(kb.clip(upper=kb.quantile(0.99)), bins=40, color="#54A24B")
    ax.set(title="Attachment size (KB, clipped at p99)", xlabel="KB", ylabel="# attachments")
    charts["attachment_size_kb"] = _save(fig, charts_dir, "attachment_size_kb"); plt.close(fig)

    # 4. extracted char counts (supported only)
    sup = att[att["supported"] & (att["charCount"] > 0)]
    if len(sup):
        cc = sup["charCount"]
        stats["extracted_chars"] = {
            "mean": int(cc.mean()), "median": int(cc.median()), "max": int(cc.max()),
            "p90": int(cc.quantile(0.9)),
        }
        fig, ax = plt.subplots()
        ax.hist(cc.clip(upper=cc.quantile(0.99)), bins=40, color="#B279A2")
        ax.set(title="Extracted characters per attachment (p99 clip)", xlabel="chars", ylabel="# attachments")
        charts["extracted_chars"] = _save(fig, charts_dir, "extracted_chars"); plt.close(fig)

    # 5. description length
    dc = tickets["descriptionChars"]
    stats["description_chars"] = {
        "mean": int(dc.mean()), "median": int(dc.median()), "max": int(dc.max()),
        "empty": int((dc == 0).sum()),
    }
    fig, ax = plt.subplots()
    ax.hist(dc.clip(upper=dc.quantile(0.99)), bins=40, color="#E45756")
    ax.set(title="Jira description length (chars, p99 clip)", xlabel="chars", ylabel="# tickets")
    charts["description_chars"] = _save(fig, charts_dir, "description_chars"); plt.close(fig)

    # 6. condense INPUT tokens vs the budget - what condense actually ingests (description +
    #    only the SELECTED attachments: idea-card alone, or top-4 supported). The all-attachments
    #    sum is reported alongside for context, but it over-counts (condense ignores the rest).
    ci = tickets["condenseInputTokens"]
    ct = tickets["combinedTokens"]
    over = int((ci > TOKEN_BUDGET).sum())
    stats["condense_input_tokens"] = {
        "budget": TOKEN_BUDGET,
        "mean": int(ci.mean()), "median": int(ci.median()), "max": int(ci.max()),
        "over_budget": over, "over_budget_pct": round(100.0 * over / max(1, len(ci)), 1),
        "all_attachments_mean_for_context": int(ct.mean()),
        "all_attachments_over_budget": int((ct > TOKEN_BUDGET).sum()),
    }
    fig, ax = plt.subplots()
    ax.hist(ci.clip(upper=max(ci.quantile(0.99), TOKEN_BUDGET * 1.1)), bins=40, color="#72B7B2",
            label="condense input (desc + selected)")
    ax.hist(ct.clip(upper=max(ct.quantile(0.99), TOKEN_BUDGET * 1.1)), bins=40, color="#F58518",
            alpha=0.35, label="all attachments (context)")
    ax.axvline(TOKEN_BUDGET, color="red", linestyle="--", label=f"{TOKEN_BUDGET:,} budget")
    ax.set(title="Tokens per ticket vs budget — what condense actually ingests", xlabel="tokens", ylabel="# tickets")
    ax.legend()
    charts["condense_input_tokens"] = _save(fig, charts_dir, "condense_input_tokens"); plt.close(fig)

    # 6b. BUDGET-APPLIED condense input - the real thing: apply the 40k/N char cap per fallback
    #     doc (idea card uncapped). This is what the LLM actually receives; tests the token thesis.
    if records:
        applied = budget_applied_input(records, budget=DOC_CHAR_BUDGET)
        ap_tokens = pd.Series([v["tokens"] for v in applied.values()])
        ap_chars = pd.Series([v["chars"] for v in applied.values()])
        paths = pd.Series([v["path"] for v in applied.values()]).value_counts()
        stats["condense_input_budget_applied"] = {
            "doc_char_budget": DOC_CHAR_BUDGET,
            "input_tokens_mean": int(ap_tokens.mean()), "input_tokens_median": int(ap_tokens.median()),
            "input_tokens_p90": int(ap_tokens.quantile(0.9)), "input_tokens_max": int(ap_tokens.max()),
            "input_chars_mean": int(ap_chars.mean()), "input_chars_max": int(ap_chars.max()),
            "tickets_over_20k_tokens": int((ap_tokens > 20_000).sum()),
            "tickets_over_30k_tokens": int((ap_tokens > 30_000).sum()),
            "by_path": {k: int(v) for k, v in paths.items()},
        }
        fig, ax = plt.subplots()
        ax.hist(ap_tokens.clip(upper=max(ap_tokens.quantile(0.99), 1)), bins=40, color="#4C78A8")
        ax.set(title=f"Condense input tokens AFTER the {DOC_CHAR_BUDGET:,}-char budget is applied",
               xlabel="tokens (input to the LLM)", ylabel="# tickets")
        charts["condense_input_budget_applied"] = _save(fig, charts_dir, "condense_input_budget_applied"); plt.close(fig)

    # 7. idea-card / supported coverage + what condense would actually select
    stats["coverage"] = {
        "tickets_with_idea_card": int(tickets["hasIdeaCard"].sum()),
        "supported_attachments": int(att["supported"].sum()),
        "unsupported_attachments": int((~att["supported"]).sum()),
        "attachments_condense_selects": int(att["selected"].sum()) if "selected" in att.columns else 0,
        "attachments_ignored": int((~att["selected"]).sum()) if "selected" in att.columns else 0,
        "extract_failures": int((att["extractError"] != "").sum()),
    }

    # ---- ticket usefulness (theme/VS ground truth) ----

    n = int(len(tickets))
    useful = tickets["themeCount"] > 0
    has_att = tickets["attachmentCount"] > 0
    uwa, udo, nu = int((useful & has_att).sum()), int((useful & ~has_att).sum()), int((~useful).sum())
    stats["usefulness"] = {
        "tickets_total": n, "useful_has_theme_gt": int(useful.sum()),
        "useful_pct": round(100.0 * int(useful.sum()) / max(1, n), 1),
        "useful_with_attachments": uwa, "useful_description_only": udo, "not_useful_no_theme": nu,
    }
    fig, ax = plt.subplots(figsize=(6, 4))
    vals = [uwa, udo, nu]
    ax.bar(["useful +\nattachments", "useful\n(desc only)", "not useful\n(no theme)"], vals,
           color=["#4C78A8", "#72B7B2", "#E45756"])
    for i, v in enumerate(vals):
        ax.text(i, v, str(v), ha="center", va="bottom")
    ax.set(title="Ticket usefulness (theme/VS ground truth present?)", ylabel="# tickets")
    charts["usefulness"] = _save(fig, charts_dir, "usefulness"); plt.close(fig)

    # themes per ticket ("2 themes -> N tickets")
    theme_dist = tickets["themeCount"].value_counts().sort_index()
    stats["themes_per_ticket"] = {f"{int(k)} themes": int(v) for k, v in theme_dist.items()}
    stats["themes_per_ticket"]["mean"] = round(float(tickets["themeCount"].mean()), 2)
    stats["themes_per_ticket"]["max"] = int(tickets["themeCount"].max())
    fig, ax = plt.subplots()
    theme_dist.plot(kind="bar", ax=ax, color="#54A24B")
    ax.set(title="Themes (Value Streams) per ticket", xlabel="# themes", ylabel="# tickets")
    charts["themes_per_ticket"] = _save(fig, charts_dir, "themes_per_ticket"); plt.close(fig)

    # themes among tickets with NO attachments
    no_att = tickets[~has_att]
    stats["no_attachment_tickets"] = {
        "count": int(len(no_att)),
        "useful_via_description": int((no_att["themeCount"] > 0).sum()),
        "dead_no_theme_no_attachment": int((no_att["themeCount"] == 0).sum()),
        **{f"{int(k)} themes": int(v) for k, v in no_att["themeCount"].value_counts().sort_index().items()},
    }
    if len(no_att):
        fig, ax = plt.subplots()
        no_att["themeCount"].value_counts().sort_index().plot(kind="bar", ax=ax, color="#B279A2")
        ax.set(title="Themes per ticket — tickets WITHOUT attachments", xlabel="# themes", ylabel="# tickets")
        charts["no_attachment_tickets"] = _save(fig, charts_dir, "no_attachment_themes"); plt.close(fig)

    # 4b. For description-only tickets (no attachments): is the fallback description enough?
    #     Their entire context is the Jira description, so its length decides whether condense
    #     has anything to work with. Bucket by thinness.
    if len(no_att):
        d = no_att["descriptionChars"]
        useful_no_att = no_att[no_att["themeCount"] > 0]["descriptionChars"]  # the ones we'd train on
        buckets = {
            "near_empty (<200)": int((d < 200).sum()),
            "thin (200-500)": int(((d >= 200) & (d < 500)).sum()),
            "ok (500-1500)": int(((d >= 500) & (d < 1500)).sum()),
            "rich (1500+)": int((d >= 1500).sum()),
        }
        stats["description_only_tickets"] = {
            "count": int(len(no_att)),
            "description_chars_mean": int(d.mean()), "description_chars_median": int(d.median()),
            "description_chars_min": int(d.min()), "description_chars_max": int(d.max()),
            "useful_and_thin_under_500": int((useful_no_att < 500).sum()),  # useful but weak context
            **buckets,
        }
        fig, ax = plt.subplots()
        ax.hist(d.clip(upper=max(d.quantile(0.99), 1)), bins=30, color="#9D755D")
        ax.axvline(500, color="orange", linestyle="--", label="thin < 500")
        ax.axvline(1500, color="green", linestyle="--", label="rich >= 1500")
        ax.set(title="Description length — tickets with NO attachments (the only context they have)",
               xlabel="description chars", ylabel="# tickets")
        ax.legend()
        charts["description_only_tickets"] = _save(fig, charts_dir, "description_only_length"); plt.close(fig)

    # Value Stream coverage
    all_vs = [vs for lst in tickets["valueStreamIds"] for vs in (lst or [])]
    name_by_id: dict[str, str] = {}
    for r in records or []:
        for i, vs in enumerate(r.get("valueStreamIds") or []):
            name_by_id.setdefault(vs, (r.get("valueStreamNames") or [vs])[i] if i < len(r.get("valueStreamNames") or []) else vs)
    vs_freq = pd.Series(all_vs).value_counts()
    stats["value_stream_coverage"] = {
        "total_theme_links": int(len(all_vs)),
        "unique_value_streams": int(pd.Series(all_vs).nunique()),
        "value_streams_seen_once": int((vs_freq == 1).sum()),
        "top_value_streams": {f"{name_by_id.get(k, k)} ({k})": int(v) for k, v in vs_freq.head(15).items()},
    }
    if len(vs_freq):
        fig, ax = plt.subplots(figsize=(8, 6))
        top = vs_freq.head(15)[::-1]
        ax.barh([str(name_by_id.get(k, k))[:40] for k in top.index], top.values, color="#4C78A8")
        ax.set(title="Top 15 Value Streams by ticket frequency", xlabel="# tickets")
        charts["value_stream_coverage"] = _save(fig, charts_dir, "value_stream_coverage"); plt.close(fig)

    return stats, charts


def collect_failures(records: list[dict]) -> list[dict]:
    """Every fetch/extract failure, for debugging: {ticketId, filename, kind, reason}."""
    out: list[dict] = []
    for r in records:
        reason = r.get("extractError") or ""
        if reason:
            out.append({
                "ticketId": r.get("ticketId", ""), "filename": r.get("filename", ""),
                "kind": r.get("kind", ""), "ext": r.get("ext", ""), "reason": reason,
            })
    return out


# ---------------------------------------------------------------- docx export

def _metric_table(doc, rows: dict) -> None:
    """A clean 2-column metric/value table with a shaded header."""
    table = doc.add_table(rows=1, cols=2)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text = "Metric", "Value"
    for k, v in rows.items():
        cells = table.add_row().cells
        cells[0].text = str(k).replace("_", " ")
        cells[1].text = f"{v:,}" if isinstance(v, int) else str(v)
    doc.add_paragraph()


def _highlights(stats: dict, n_tickets: int) -> dict:
    """The headline numbers, pulled from the per-section stats for the summary table."""
    pres = stats.get("attachment_presence", {})
    size = stats.get("attachment_size_kb", {})
    ct = stats.get("condense_input_tokens", {})
    cov = stats.get("coverage", {})
    use = stats.get("usefulness", {})
    vc = stats.get("value_stream_coverage", {})
    return {
        "Tickets analyzed": n_tickets,
        "Useful (has theme/VS GT)": f"{use.get('useful_has_theme_gt', '?')} ({use.get('useful_pct', '?')}%)",
        "Useful with attachments": use.get("useful_with_attachments", "?"),
        "Useful (description only)": use.get("useful_description_only", "?"),
        "Not useful (no theme)": use.get("not_useful_no_theme", "?"),
        "Tickets without attachments": f"{pres.get('tickets_without_attachments', '?')} "
                                       f"({pres.get('pct_without', '?')}%)",
        "Most common file type": stats.get("most_common_type", "?"),
        "Avg attachment size (KB)": size.get("per_attachment_mean", "?"),
        "Tickets over budget (condense input)": f"{ct.get('over_budget', '?')} ({ct.get('over_budget_pct', '?')}%)",
        "Tickets with an idea card": cov.get("tickets_with_idea_card", "?"),
        "Unique Value Streams covered": vc.get("unique_value_streams", "?"),
    }


def build_docx(stats: dict, charts: dict, out_path: Path, *, n_tickets: int,
               failures: list[dict] | None = None) -> None:
    from datetime import date

    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt, RGBColor

    doc = Document()
    title = doc.add_heading("IDMT Ticket EDA — Phase 1", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph(f"Usefulness + attachments · {n_tickets} tickets · {date.today().isoformat()}")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in sub.runs:
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_heading("Highlights", level=1)
    doc.add_paragraph(
        "Profile of the historical IDMT tickets: how usable they are (theme/VS ground truth) and "
        "their attachment shape, which drives condense/ingestion sizing."
    )
    _metric_table(doc, _highlights(stats, n_tickets))

    # title, stat key, one-line explanation of what the chart answers
    sections = [
        ("1. Ticket usefulness", "usefulness",
         "Useful = has >=1 theme with a Value Stream (ground truth), split by whether it also has attachments."),
        ("2. Themes per ticket", "themes_per_ticket",
         "How many Value Streams each ticket maps to (the '2 themes -> N tickets' distribution)."),
        ("3. Themes among tickets without attachments", "no_attachment_tickets",
         "For tickets with no attachment: how many still carry theme ground truth (usable via description)."),
        ("3b. Description-only context — is the fallback enough?", "description_only_tickets",
         "Tickets with no attachments fall back to the Jira description alone. This shows how thin that "
         "context is (near_empty/thin/ok/rich), and how many useful tickets have a weak (<500 char) "
         "description - the cases where description-only may not be enough for prediction/generation."),
        ("4. Value Stream coverage", "value_stream_coverage",
         "Unique Value Streams covered and which dominate; ones seen once are thin evidence."),
        ("5. Tickets with vs without attachments", "attachment_presence",
         "How many tickets carry no attachment at all — those fall back to the Jira description."),
        ("6. Attachments per ticket", "attachments_per_ticket",
         "Distribution of attachment count per ticket; informs the top-N selection cap."),
        ("7. File types", "file_types",
         "What attachment formats appear and which dominates (idea cards are usually PPT/PPTX)."),
        ("8. Attachment size", "attachment_size_kb",
         "Per-attachment size and the total payload per ticket; informs the pre-download size cap."),
        ("9. Extracted characters per attachment", "extracted_chars",
         "How much text each supported attachment yields after extraction."),
        ("10. Jira description length", "description_chars",
         "Typical Jira description length — the fallback context when no idea card exists."),
        ("11. Condense input tokens vs budget", "condense_input_tokens",
         "Tokens condense actually ingests = description + only the SELECTED attachments (idea card "
         "alone, or top-4 supported). The all-attachments sum is shown for context but over-counts, "
         "since condense never feeds every attachment. NOTE: this uses full extracted text - the next "
         "section applies the real per-doc char cap."),
        ("11b. Condense input AFTER the char budget", "condense_input_budget_applied",
         "The true input to the LLM: idea-card path uncapped, fallback path caps each selected doc at "
         "doc_char_budget/N chars. This is the number that tests the token threshold thesis."),
        ("12. Coverage", "coverage",
         "Idea-card presence, supported vs unsupported attachments, condense selection, extraction failures."),
    ]
    doc.add_page_break()
    for title_text, key, blurb in sections:
        if key not in stats and key not in charts:
            continue
        doc.add_heading(title_text, level=1)
        doc.add_paragraph(blurb)
        if key in charts and Path(charts[key]).exists():
            doc.add_picture(charts[key], width=Inches(6.0))
        if key in stats and isinstance(stats[key], dict):
            flat = {k: v for k, v in stats[key].items() if not isinstance(v, (dict, list))}
            if flat:
                _metric_table(doc, flat)
            for sub_name, sub_rows in stats[key].items():
                if isinstance(sub_rows, dict) and sub_rows:
                    doc.add_paragraph(sub_name.replace("_", " ") + ":")
                    _metric_table(doc, sub_rows)

    # 9. Failures - ticket id + attachment + reason, for debugging
    doc.add_heading("13. Fetch / extract failures", level=1)
    if failures:
        doc.add_paragraph(
            f"{len(failures)} attachment(s)/ticket(s) failed to fetch or extract (timeout, "
            "download error, or unparseable file). Listed for debugging."
        )
        table = doc.add_table(rows=1, cols=3)
        table.style = "Light Grid Accent 1"
        for cell, label in zip(table.rows[0].cells, ("Ticket", "Attachment", "Reason")):
            cell.text = label
        for f in failures:
            cells = table.add_row().cells
            cells[0].text = str(f.get("ticketId", ""))
            cells[1].text = str(f.get("filename", ""))
            cells[2].text = str(f.get("reason", ""))[:200]
    else:
        doc.add_paragraph("No fetch or extract failures.")

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))


# ---------------------------------------------------------------- CLI

async def _main(args) -> None:
    out = Path(args.out)
    out.mkdir(parents=True, exist_ok=True)
    cache = out / "attachments_raw.json"

    ticket_ids = load_ticket_ids(args.tickets)
    have = json.loads(cache.read_text(encoding="utf-8")) if cache.exists() else []
    # Only count a ticket as cached if it has theme data (new schema); old attachments-only
    # caches lack themeCount and would show every ticket as 0 themes / not useful.
    collected_ids = {r["ticketId"] for r in have
                     if r.get("kind") == "description" and "themeCount" in r}
    remaining = [t for t in ticket_ids if t not in collected_ids]

    if args.use_cache and not remaining:
        records = have
        print(f"loaded {len(records)} cached records ({len(collected_ids)} tickets) from {cache}")
    else:
        # Discover the Business Value Stream field id (for theme/usefulness counting).
        settings = load_settings()
        async with httpx.AsyncClient(
            base_url=settings.jira_base_url,
            headers={"Authorization": f"Bearer {settings.jira_token}"},
            timeout=settings.jira_timeout_seconds, verify=settings.jira_verify_ssl,
        ) as http:
            bvs_field = args.field_id or await discover_field_id(http, settings.jira_api_version, args.field_name)
        if not bvs_field:
            raise SystemExit(f"could not find '{args.field_name}' field; pass --field-id customfield_#####")
        print(f"Business Value Stream field id: {bvs_field}")
        # Always resume from whatever is already cached, checkpointing as we go, so a stall/stop
        # never loses progress - re-run the same command to continue from where it stopped.
        print(f"collecting {len(remaining)}/{len(ticket_ids)} tickets "
              f"(concurrency={args.concurrency}, timeout={args.ticket_timeout}s, "
              f"checkpoint every {args.checkpoint_every}); cache: {cache}")
        records = await collect(
            ticket_ids, bvs_field=bvs_field, concurrency=args.concurrency, ticket_timeout=args.ticket_timeout,
            done_records=have, checkpoint_path=cache, checkpoint_every=args.checkpoint_every,
        )
        print(f"collected {len({r['ticketId'] for r in records})} tickets -> {cache}")

    att, tickets = to_frames(records)
    stats, charts = build_charts(att, tickets, out / "charts", records=records)
    (out / "stats.json").write_text(json.dumps(stats, indent=2), encoding="utf-8")
    print(f"\nstats -> {out/'stats.json'}\ncharts -> {out/'charts'}/")
    print(json.dumps(stats, indent=2))

    # Failures (ticket id + attachment + reason) for debugging - own file + console.
    failures = collect_failures(records)
    (out / "failures.json").write_text(json.dumps(failures, indent=2), encoding="utf-8")
    print(f"\n{len(failures)} fetch/extract failures -> {out/'failures.json'}")
    for f in failures[:30]:
        print(f"  {f['ticketId']}  {f['filename']}  -> {f['reason']}")
    if len(failures) > 30:
        print(f"  ... and {len(failures) - 30} more (see failures.json)")

    if args.docx:
        docx_path = out / "attachments_eda.docx"
        build_docx(stats, charts, docx_path, n_tickets=int(len(tickets)), failures=failures)
        print(f"docx -> {docx_path}")


if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("tickets", help="text file with one IDMT id per line")
    parser.add_argument("--out", default="out/eda/attachments")
    parser.add_argument("--field-name", default="Business Value Stream",
                        help="display name of the theme Value Stream field (for usefulness)")
    parser.add_argument("--field-id", default="", help="explicit customfield_##### (skips discovery)")
    parser.add_argument("--concurrency", type=int, default=4)
    parser.add_argument("--ticket-timeout", type=float, default=90.0,
                        help="seconds before a slow fetch/download/extract is abandoned (per call)")
    parser.add_argument("--checkpoint-every", type=int, default=10,
                        help="flush the cache every N tickets so a stop never loses progress")
    parser.add_argument("--use-cache", action="store_true",
                        help="if every ticket is already cached, skip collection; otherwise resume")
    parser.add_argument("--docx", action="store_true", help="also write the .docx report")
    asyncio.run(_main(parser.parse_args()))
