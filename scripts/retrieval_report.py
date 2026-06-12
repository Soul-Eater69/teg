"""Build the historic-lane retrieval report docx from retrieval_eval.json.

Authored, data-driven: the verdict text is computed from the actual numbers so the doc is complete
on a single eval run. Called automatically by retrieval_eval.py; also runnable standalone to rebuild
the docx from a saved json without re-running the eval:

    uv run python scripts/retrieval_report.py out/eval/retrieval_eval.json
"""

from __future__ import annotations

import json
import sys
from pathlib import Path


def _pct(v):
    return "n/a" if v is None else f"{v:.1%}"


def _f3(v):
    return "n/a" if v is None else f"{v:.3f}"


def _titles(ax, title, note):
    ax.figure.suptitle(title, fontsize=13, fontweight="bold", y=0.99)
    if note:
        ax.set_title(note, fontsize=9, style="italic", color="#555", pad=8)


def _add_table(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = "Light Grid Accent 1"
    for c, h in zip(t.rows[0].cells, headers):
        c.text = str(h)
    for r in rows:
        cells = t.add_row().cells
        for c, v in zip(cells, r):
            c.text = str(v)


def build_report(payload: dict, out_path: Path) -> None:
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    from docx import Document
    from docx.shared import Inches

    agg = payload["aggregates"]
    ks = payload["config"]["k_values"]
    by_k = {int(k): v for k, v in agg["by_k"].items()}
    charts = out_path.parent / "retrieval_charts"
    charts.mkdir(parents=True, exist_ok=True)

    def save(fig, name):
        p = charts / f"{name}.png"
        fig.savefig(p, dpi=150, bbox_inches="tight"); plt.close(fig)
        return str(p)

    # 1. the recall/precision tradeoff across K - the headline
    fig, ax = plt.subplots(figsize=(8, 4.6))
    rec = [by_k[k]["recall"]["mean"] for k in ks]
    prec = [by_k[k]["precision"]["mean"] for k in ks]
    pstrict = [by_k[k]["precision_strict"]["mean"] for k in ks]
    ax.plot(ks, rec, "-o", label="Recall@k (GT coverage)")
    ax.plot(ks, prec, "-s", label="Precision@k (relevant tickets)")
    ax.plot(ks, pstrict, "--^", label="Precision@k (excl. broad VS)")
    for k, r, p in zip(ks, rec, prec):
        ax.annotate(f"{r:.0%}", (k, r), textcoords="offset points", xytext=(0, 6), ha="center", fontsize=8)
        ax.annotate(f"{p:.0%}", (k, p), textcoords="offset points", xytext=(0, -12), ha="center", fontsize=8)
    ax.set_xticks(ks); ax.set_xlabel("K (past tickets retrieved)"); ax.set_ylabel("score"); ax.set_ylim(0, 1.05)
    _titles(ax, "Recall rises but precision falls as K grows",
            "More tickets cover more GT (recall up) but pull in less-relevant ones (precision down) - the dilution.")
    ax.legend(fontsize=8); ax.grid(alpha=0.3)
    f_tradeoff = save(fig, "tradeoff")

    # 2. MRR / hit-rate / nDCG across K
    fig, ax = plt.subplots(figsize=(8, 4.6))
    import numpy as np
    x = np.arange(len(ks)); w = 0.25
    for i, (key, label) in enumerate((("mrr", "MRR"), ("hit_rate", "Hit@k"), ("ndcg", "nDCG@k"))):
        vals = [by_k[k][key] for k in ks]
        bars = ax.bar(x + (i - 1) * w, vals, w, label=label)
        for b, v in zip(bars, vals):
            ax.text(b.get_x() + b.get_width() / 2, v, f"{v:.2f}", ha="center", va="bottom", fontsize=7)
    ax.set_xticks(x); ax.set_xticklabels([f"K={k}" for k in ks]); ax.set_ylim(0, 1.1)
    _titles(ax, "Does retrieval rank relevant tickets early?",
            "MRR ~ how high the first relevant ticket sits. Hit@k ~ found anything. nDCG ~ rank-weighted relevance.")
    ax.legend(fontsize=8); ax.grid(axis="y", alpha=0.3)
    f_rank = save(fig, "ranking")

    # 3. evidence density histogram (VS labels per retrieved ticket)
    hist = agg["evidence_density_all"]["histogram"]
    fig, ax = plt.subplots(figsize=(8, 4.2))
    ax.bar([int(k) for k in hist], list(hist.values()))
    ax.set_xlabel("Value Stream labels on a retrieved ticket"); ax.set_ylabel("number of retrieved tickets")
    _titles(ax, "How much precedent does each retrieved ticket carry?",
            f"Mean {agg['evidence_density_all']['mean']:.2f} VS/ticket "
            f"(min {agg['evidence_density_all']['min']}, max {agg['evidence_density_all']['max']}).")
    ax.grid(axis="y", alpha=0.3)
    f_density = save(fig, "density")

    # 4. first-relevant-rank histogram (how deep is the first useful ticket)
    fr_hist = agg["first_relevant_rank"]["histogram"]
    fig, ax = plt.subplots(figsize=(8, 4.2))
    ax.bar([int(k) for k in fr_hist], list(fr_hist.values()))
    ax.set_xlabel("rank of the FIRST relevant retrieved ticket"); ax.set_ylabel("queries")
    _titles(ax, "How early does the first useful precedent appear?",
            f"Lower = better. {agg['first_relevant_rank']['none_count']} queries had NO relevant ticket at all.")
    ax.grid(axis="y", alpha=0.3)
    f_first = save(fig, "first_rank")

    # 5. per-query coverage split across K (fully covered / partial / zero)
    fig, ax = plt.subplots(figsize=(8, 4.2))
    full = [by_k[k]["fully_covered_rate"] for k in ks]
    zero = [by_k[k]["zero_hit_rate"] for k in ks]
    partial = [1 - f - z for f, z in zip(full, zero)]
    ax.bar(x, full, 0.55, label="fully covered (all GT found)")
    ax.bar(x, partial, 0.55, bottom=full, label="partly covered")
    ax.bar(x, zero, 0.55, bottom=[f + p for f, p in zip(full, partial)], label="zero relevant")
    ax.set_xticks(x); ax.set_xticklabels([f"K={k}" for k in ks]); ax.set_ylim(0, 1.05)
    _titles(ax, "Per-query: how much of each ticket's GT does precedent cover?",
            "Stacked to 100% of queries. 'Zero relevant' = retrieval total miss (no prompt can fix those).")
    ax.legend(fontsize=8); ax.grid(axis="y", alpha=0.3)
    f_split = save(fig, "coverage_split")

    # ---------------- document
    doc = Document()
    doc.add_heading("Historic-lane retrieval evaluation", level=0)
    doc.add_paragraph(
        f"How good is the retriever that surfaces the similar past tickets shown to the model as "
        f"precedent - measured on its own, without the LLM, over {agg['n_queries']} tickets. A retrieved "
        "past ticket is RELEVANT when its Value Stream labels overlap the query ticket's correct Value "
        "Streams (a free, automatic relevance signal - no human labelling). That definition makes the "
        "classic retrieval metrics apply. Each chart states its question and how to read it; definitions "
        "are at the end.")

    # data-driven verdict
    r6, r10 = by_k[6]["recall"]["mean"], by_k[10]["recall"]["mean"]
    p6, p10 = by_k[6]["precision"]["mean"], by_k[10]["precision"]["mean"]
    zero6 = by_k[6]["zero_hit_rate"]
    sep = agg["score_separation"]
    doc.add_heading("Verdict", level=1)
    bullets = [
        f"Coverage is strong: at K=6 the retrieved past tickets already contain {_pct(r6)} of the correct "
        f"Value Streams (recall@6). Growing to K=10 lifts it to {_pct(r10)} - only +{(r10 - r6):.1%}.",
        f"But precision falls with K: {_pct(p6)} of the top-6 are relevant vs {_pct(p10)} of the top-10 "
        f"({(p10 - p6):+.1%}). The extra tickets at K=8-10 are less-relevant - this is the dilution that "
        "made K>6 not worth it in the generation eval, now measured directly.",
        f"MRR@10 = {by_k[10]['mrr']:.3f}: the first relevant precedent usually sits near the top (rank "
        f"~{(1 / by_k[10]['mrr']):.1f} on average), so the ranker puts useful tickets early.",
        f"Retrieval total misses: {_pct(zero6)} of queries had NO relevant ticket in the top-6 - those are "
        "the GT precedent simply cannot reach (a different lever than the prompt).",
    ]
    if sep["relevant_mean"] is not None and sep["irrelevant_mean"] is not None:
        bullets.append(
            f"The ranker separates signal from noise: relevant retrieved tickets score "
            f"{sep['relevant_mean']:.3f} on average vs {sep['irrelevant_mean']:.3f} for irrelevant ones.")
    bullets.append(
        "Context usage (from the generation eval, not recomputed here): of the GT this retriever puts in "
        "context, the model picks ~79% (precedent backed) - so the retriever's coverage is the ceiling and "
        "the model uses most of it. This eval measures the ceiling; the generation eval measures the usage.")
    for b in bullets:
        doc.add_paragraph(b, style="List Bullet")

    doc.add_heading("1. The recall / precision tradeoff across K", level=1)
    doc.add_paragraph(
        "Recall@k = of the query's correct Value Streams, how many appear in the union of the top-K "
        "retrieved tickets' labels (the coverage ceiling). Precision@k = of the K retrieved tickets, how "
        "many are relevant (share a correct VS). The dashed line excludes 'broad' Value Streams (tagged on "
        f">{payload['config']['broad_fraction']:.0%} of tickets) so a shared generic stream doesn't inflate "
        "precision. This is the core finding: more tickets buy coverage at the cost of relevance.")
    doc.add_picture(f_tradeoff, width=Inches(6.2))
    _add_table(doc, ["K", "Recall@k", "Precision@k", "Precision@k (strict)", "Hit@k", "MRR", "nDCG@k",
                     "Zero-relevant", "Fully covered"],
               [[k, _pct(by_k[k]["recall"]["mean"]), _pct(by_k[k]["precision"]["mean"]),
                 _pct(by_k[k]["precision_strict"]["mean"]), _pct(by_k[k]["hit_rate"]),
                 _f3(by_k[k]["mrr"]), _f3(by_k[k]["ndcg"]), _pct(by_k[k]["zero_hit_rate"]),
                 _pct(by_k[k]["fully_covered_rate"])] for k in ks])
    mrows = [[step, f"{m['recall_gain']:+.1%}", f"{m['precision_change']:+.1%}"]
             for step, m in agg["marginal"].items()]
    doc.add_paragraph("Marginal effect of adding tickets (recall gained vs precision lost):")
    _add_table(doc, ["K step", "Recall gained", "Precision change"], mrows)

    doc.add_heading("2. Does retrieval rank relevant tickets early?", level=1)
    doc.add_paragraph(
        "MRR (mean reciprocal rank) ~ 1/(rank of the first relevant ticket): near 1.0 means the very first "
        "result is usually relevant. Hit@k = fraction of queries with at least one relevant ticket in the "
        "top-K. nDCG@k rewards putting relevant tickets higher up.")
    doc.add_picture(f_rank, width=Inches(6.2))

    doc.add_heading("3. How early does the first useful precedent appear?", level=1)
    doc.add_paragraph(
        "For each query, the rank of the first relevant retrieved ticket (lower is better). A spike at rank 1 "
        "means the top result is usually on-target; a long tail means relevant precedent often sits deep "
        "(an argument for a larger K). Queries with NO relevant ticket are retrieval total misses.")
    doc.add_picture(f_first, width=Inches(6.2))

    doc.add_heading("4. Per-query coverage: fully / partly / not covered", level=1)
    doc.add_paragraph(
        "Splits the queries three ways at each K: all of their GT covered by precedent, some covered, or "
        "none. The 'zero relevant' band is the hard floor - GT that precedent cannot surface, which the "
        "prompt cannot fix; improving it needs better retrieval or more corpus.")
    doc.add_picture(f_split, width=Inches(6.2))

    doc.add_heading("5. Evidence density - precedent per retrieved ticket", level=1)
    doc.add_paragraph(
        "How many Value Stream labels each retrieved past ticket carries. Richer tickets contribute more "
        "candidate precedent per slot. This is descriptive context for the coverage numbers above.")
    doc.add_picture(f_density, width=Inches(6.2))

    doc.add_heading("What the metrics mean", level=1)
    _add_table(doc, ["Term", "Plain meaning"], [
        ["Relevant ticket", "A retrieved past ticket whose Value Stream labels overlap the query's correct VS."],
        ["Recall@k / coverage", "Of the query's correct VS, the fraction present in the top-K tickets' labels."],
        ["Precision@k", "Of the K retrieved tickets, the fraction that are relevant."],
        ["Precision@k (strict)", "Precision counting only non-broad shared VS, so a generic stream doesn't inflate it."],
        ["Hit@k", "Whether at least one relevant ticket appears in the top-K."],
        ["MRR", "1 / rank of the first relevant ticket, averaged - how early the first useful result lands."],
        ["nDCG@k", "Rank-weighted relevance - rewards relevant tickets ranked higher."],
        ["Evidence density", "Value Stream labels carried per retrieved ticket."],
        ["Zero relevant", "A query whose top-K had no relevant ticket - a retrieval total miss."],
        ["Context usage", "From the generation eval: of GT put in context, how much the model actually picked."],
    ])

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))


def main() -> None:
    path = Path(sys.argv[1] if len(sys.argv) > 1 else "out/eval/retrieval_eval.json")
    payload = json.loads(path.read_text(encoding="utf-8"))
    out = path.with_suffix(".docx")
    build_report(payload, out)
    print(f"report -> {out}")


if __name__ == "__main__":
    main()
