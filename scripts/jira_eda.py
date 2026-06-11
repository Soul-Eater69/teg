"""Phase 1 EDA — Jira ticket usefulness across the historical IDMT set.

Metadata only (no attachment download/extract, no condense) so it is fast over the full set.
Per ticket it counts attachments and linked themes (a theme = a linked issue, ANY link type,
that carries a Business Value Stream value), then profiles how many tickets are actually usable
for training: a ticket is USEFUL when it has >= 1 theme (Value Stream ground truth).

Answers: attachments per ticket, themes per ticket (the "2 themes -> N tickets" distribution),
the theme distribution among tickets with NO attachments, the usefulness split (useful with
attachments / useful description-only / not useful), and Value Stream coverage.

Usage (needs Jira creds in .env; install: uv sync --extra eda):
  uv run python scripts/jira_eda.py tickets.txt --out out/eda/jira --docx
  uv run python scripts/jira_eda.py tickets.txt --out out/eda/jira --docx   # re-run resumes
"""

from __future__ import annotations

import argparse
import asyncio
import json
from pathlib import Path

import httpx

from teg.condense.attachment_ranker import is_idea_card, is_supported
from teg.config.settings import load_settings
from teg.ingestion.extraction.value_stream_field import parse_value_stream


def load_ticket_ids(path: str) -> list[str]:
    lines = Path(path).read_text(encoding="utf-8").splitlines()
    return [s.strip() for s in lines if s.strip() and not s.strip().startswith("#")]


# ---------------------------------------------------------------- collection

async def _get(http, api, key, fields) -> dict:
    resp = await http.get(f"/rest/api/{api}/issue/{key}", params={"fields": fields})
    resp.raise_for_status()
    return resp.json() or {}


async def _discover_field_id(http, api, field_name) -> str | None:
    resp = await http.get(f"/rest/api/{api}/field")
    resp.raise_for_status()
    wanted = field_name.strip().lower()
    for f in resp.json() or []:
        if str(f.get("name") or "").strip().lower() == wanted:
            return str(f.get("id") or "")
    return None


def _linked_keys(fields: dict) -> list[str]:
    keys: list[str] = []
    for link in fields.get("issuelinks") or []:
        issue = link.get("inwardIssue") or link.get("outwardIssue")
        if isinstance(issue, dict):
            key = str(issue.get("key") or "")
            if key and key not in keys:
                keys.append(key)
    return keys


async def collect(
    ticket_ids, *, bvs_field, concurrency=6, ticket_timeout=60.0,
    done_records=None, checkpoint_path=None, checkpoint_every=20,
) -> list[dict]:
    """One record per ticket: counts + the value streams its themes carry. Resumable + timed out."""
    settings = load_settings()
    http = httpx.AsyncClient(
        base_url=settings.jira_base_url,
        headers={"Authorization": f"Bearer {settings.jira_token}"},
        timeout=settings.jira_timeout_seconds, verify=settings.jira_verify_ssl,
    )
    api = settings.jira_api_version
    records: list[dict] = list(done_records or [])
    done = {r["ticketId"] for r in records}
    todo = [t for t in ticket_ids if t not in done]
    progress = {"n": 0, "total": len(todo)}

    async def _one(ticket_id: str) -> dict:
        try:
            issue = await asyncio.wait_for(
                _get(http, api, ticket_id, "summary,description,attachment,issuelinks"),
                timeout=ticket_timeout,
            )
        except Exception as exc:
            return {"ticketId": ticket_id, "error": f"{type(exc).__name__}: {exc}",
                    "attachmentCount": 0, "supportedAttachmentCount": 0, "hasIdeaCard": False,
                    "descriptionChars": 0, "linkCount": 0, "themeCount": 0, "valueStreamIds": []}
        f = issue.get("fields") or {}
        atts = f.get("attachment") or []
        linked_keys = _linked_keys(f)

        async def _theme_vs(key: str):
            try:
                li = await asyncio.wait_for(_get(http, api, key, f"summary,{bvs_field}"), timeout=ticket_timeout)
            except Exception:
                return None
            return parse_value_stream((li.get("fields") or {}).get(bvs_field))

        vs_parsed = await asyncio.gather(*(_theme_vs(k) for k in linked_keys)) if linked_keys else []
        vs_ids = [p[1] for p in vs_parsed if p]
        vs_names = [p[0] for p in vs_parsed if p]
        return {
            "ticketId": ticket_id,
            "ticketStableId": str(issue.get("id") or ""),
            "title": str(f.get("summary") or ""),
            "attachmentCount": len(atts),
            "supportedAttachmentCount": sum(1 for a in atts if is_supported(str(a.get("filename") or ""))),
            "hasIdeaCard": any(is_idea_card(str(a.get("filename") or "")) for a in atts),
            "descriptionChars": len(str(f.get("description") or "")),
            "linkCount": len(linked_keys),
            "themeCount": len(vs_ids),  # links that carry a Value Stream
            "valueStreamIds": vs_ids,
            "valueStreamNames": vs_names,
        }

    sem = asyncio.Semaphore(concurrency)

    async def _guarded(ticket_id: str) -> None:
        async with sem:
            rec = await _one(ticket_id)
        records.append(rec)
        progress["n"] += 1
        print(f"[{progress['n']}/{progress['total']}] {ticket_id}: "
              f"{rec['attachmentCount']} attach, {rec['themeCount']} themes"
              + (f"  ERROR {rec['error']}" if rec.get("error") else ""))
        if checkpoint_path is not None and progress["n"] % checkpoint_every == 0:
            checkpoint_path.write_text(json.dumps(records, ensure_ascii=False, indent=2), encoding="utf-8")

    try:
        if done:
            print(f"resuming: {len(done)} tickets already collected, {len(todo)} to go")
        await asyncio.gather(*(_guarded(t) for t in todo))
    finally:
        await http.aclose()
    if checkpoint_path is not None:
        checkpoint_path.write_text(json.dumps(records, ensure_ascii=False, indent=2), encoding="utf-8")
    return records


# ---------------------------------------------------------------- analysis

def _save(fig, charts_dir: Path, name: str) -> str:
    charts_dir.mkdir(parents=True, exist_ok=True)
    path = charts_dir / f"{name}.png"
    fig.savefig(path, dpi=150, bbox_inches="tight")
    return str(path)


def analyze(records: list[dict], charts_dir: Path) -> tuple[dict, dict]:
    import matplotlib

    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    import pandas as pd

    plt.rcParams.update({"figure.figsize": (8, 4.5), "axes.grid": True, "grid.alpha": 0.3})
    df = pd.DataFrame(records)
    n = len(df)
    stats: dict = {}
    charts: dict = {}

    useful = df["themeCount"] > 0  # has Value Stream ground truth
    has_att = df["attachmentCount"] > 0

    # 1. usefulness split
    useful_with_att = int((useful & has_att).sum())
    useful_desc_only = int((useful & ~has_att).sum())
    not_useful = int((~useful).sum())
    stats["usefulness"] = {
        "tickets_total": n,
        "useful_has_theme_gt": int(useful.sum()),
        "useful_pct": round(100.0 * int(useful.sum()) / max(1, n), 1),
        "useful_with_attachments": useful_with_att,
        "useful_description_only": useful_desc_only,
        "not_useful_no_theme": not_useful,
    }
    fig, ax = plt.subplots(figsize=(6, 4))
    labels = ["useful +\nattachments", "useful\n(desc only)", "not useful\n(no theme)"]
    vals = [useful_with_att, useful_desc_only, not_useful]
    ax.bar(labels, vals, color=["#4C78A8", "#72B7B2", "#E45756"])
    for i, v in enumerate(vals):
        ax.text(i, v, str(v), ha="center", va="bottom")
    ax.set(title="Ticket usefulness (theme/VS ground truth present?)", ylabel="# tickets")
    charts["usefulness"] = _save(fig, charts_dir, "usefulness"); plt.close(fig)

    # 2. themes per ticket distribution ("2 themes -> N tickets")
    theme_dist = df["themeCount"].value_counts().sort_index()
    stats["themes_per_ticket"] = {f"{k} themes": int(v) for k, v in theme_dist.items()}
    stats["themes_per_ticket"]["mean"] = round(float(df["themeCount"].mean()), 2)
    stats["themes_per_ticket"]["max"] = int(df["themeCount"].max())
    fig, ax = plt.subplots()
    theme_dist.plot(kind="bar", ax=ax, color="#54A24B")
    ax.set(title="Themes (Value Streams) per ticket", xlabel="# themes", ylabel="# tickets")
    charts["themes_per_ticket"] = _save(fig, charts_dir, "themes_per_ticket"); plt.close(fig)

    # 3. attachments per ticket distribution
    att_dist = df["attachmentCount"].value_counts().sort_index()
    stats["attachments_per_ticket"] = {f"{k} attach": int(v) for k, v in att_dist.items()}
    stats["attachments_per_ticket"]["mean"] = round(float(df["attachmentCount"].mean()), 2)
    stats["attachments_per_ticket"]["tickets_without_any"] = int((df["attachmentCount"] == 0).sum())
    fig, ax = plt.subplots()
    att_dist.plot(kind="bar", ax=ax, color="#F58518")
    ax.set(title="Attachments per ticket", xlabel="# attachments", ylabel="# tickets")
    charts["attachments_per_ticket"] = _save(fig, charts_dir, "attachments_per_ticket"); plt.close(fig)

    # 4. for tickets with NO attachments: how many themes do they have
    no_att = df[~has_att]
    no_att_theme_dist = no_att["themeCount"].value_counts().sort_index()
    stats["no_attachment_tickets"] = {
        "count": int(len(no_att)),
        "of_those_useful_have_theme": int((no_att["themeCount"] > 0).sum()),
        "of_those_no_theme_and_no_attachment": int((no_att["themeCount"] == 0).sum()),
        **{f"{k} themes": int(v) for k, v in no_att_theme_dist.items()},
    }
    if len(no_att):
        fig, ax = plt.subplots()
        no_att_theme_dist.plot(kind="bar", ax=ax, color="#B279A2")
        ax.set(title="Themes per ticket — tickets WITHOUT attachments", xlabel="# themes", ylabel="# tickets")
        charts["no_attachment_tickets"] = _save(fig, charts_dir, "no_attachment_themes"); plt.close(fig)

    # 5. Value Stream coverage
    all_vs = [vs for r in records for vs in (r.get("valueStreamIds") or [])]
    vs_series = pd.Series(all_vs)
    vs_freq = vs_series.value_counts()
    name_by_id = {}
    for r in records:
        for i, vs in enumerate(r.get("valueStreamIds") or []):
            name_by_id.setdefault(vs, (r.get("valueStreamNames") or [None] * (i + 1))[i])
    stats["value_stream_coverage"] = {
        "total_theme_links": int(len(all_vs)),
        "unique_value_streams": int(vs_series.nunique()),
        "value_streams_seen_once": int((vs_freq == 1).sum()),
        "top_value_streams": {f"{name_by_id.get(k, k)} ({k})": int(v) for k, v in vs_freq.head(15).items()},
    }
    if len(vs_freq):
        fig, ax = plt.subplots(figsize=(8, 6))
        top = vs_freq.head(15)[::-1]
        ax.barh([str(name_by_id.get(k, k))[:40] for k in top.index], top.values, color="#4C78A8")
        ax.set(title="Top 15 Value Streams by ticket frequency", xlabel="# tickets")
        charts["value_stream_coverage"] = _save(fig, charts_dir, "value_stream_coverage"); plt.close(fig)

    # 6. errors
    errs = [r for r in records if r.get("error")]
    stats["errors"] = {"count": len(errs), "tickets": [r["ticketId"] for r in errs][:50]}
    return stats, charts


# ---------------------------------------------------------------- docx

def _metric_table(doc, rows: dict) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.style = "Light Grid Accent 1"
    table.rows[0].cells[0].text, table.rows[0].cells[1].text = "Metric", "Value"
    for k, v in rows.items():
        cells = table.add_row().cells
        cells[0].text = str(k).replace("_", " ")
        cells[1].text = f"{v:,}" if isinstance(v, int) else str(v)
    doc.add_paragraph()


def build_docx(stats: dict, charts: dict, out_path: Path, *, n_tickets: int) -> None:
    from datetime import date

    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt, RGBColor

    doc = Document()
    t = doc.add_heading("Jira Ticket EDA — Phase 1", level=0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph(f"IDMT usefulness · {n_tickets} tickets · {date.today().isoformat()}")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in sub.runs:
        run.font.size = Pt(11); run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    u = stats.get("usefulness", {})
    vc = stats.get("value_stream_coverage", {})
    doc.add_heading("Highlights", level=1)
    doc.add_paragraph(
        "How many of the historical tickets are actually usable for training. A ticket is "
        "useful when it carries at least one theme with a Value Stream (ground truth)."
    )
    _metric_table(doc, {
        "Tickets analyzed": n_tickets,
        "Useful (has theme/VS GT)": f"{u.get('useful_has_theme_gt', '?')} ({u.get('useful_pct', '?')}%)",
        "Useful with attachments": u.get("useful_with_attachments", "?"),
        "Useful (description only)": u.get("useful_description_only", "?"),
        "Not useful (no theme)": u.get("not_useful_no_theme", "?"),
        "Unique Value Streams covered": vc.get("unique_value_streams", "?"),
        "Value Streams seen once": vc.get("value_streams_seen_once", "?"),
    })

    sections = [
        ("1. Ticket usefulness", "usefulness",
         "Useful = has >=1 theme with a Value Stream. Split by whether it also has attachments."),
        ("2. Themes per ticket", "themes_per_ticket",
         "How many Value Streams each ticket maps to (the '2 themes -> N tickets' distribution)."),
        ("3. Attachments per ticket", "attachments_per_ticket",
         "Attachment count distribution; tickets_without_any fall back to the description."),
        ("4. Themes among tickets without attachments", "no_attachment_tickets",
         "For tickets with no attachment: how many still carry theme ground truth (usable via description)."),
        ("5. Value Stream coverage", "value_stream_coverage",
         "How many unique Value Streams appear and which dominate; ones seen once are thin evidence."),
        ("6. Errors", "errors", "Tickets that failed to fetch (timeout or API error)."),
    ]
    doc.add_page_break()
    for title_text, key, blurb in sections:
        doc.add_heading(title_text, level=1)
        doc.add_paragraph(blurb)
        if key in charts and Path(charts[key]).exists():
            doc.add_picture(charts[key], width=Inches(6.0))
        if key in stats and isinstance(stats[key], dict):
            flat = {k: v for k, v in stats[key].items() if not isinstance(v, (dict, list))}
            if flat:
                _metric_table(doc, flat)
            nested = {k: v for k, v in stats[key].items() if isinstance(v, dict)}
            for sub_name, sub_rows in nested.items():
                doc.add_paragraph(sub_name.replace("_", " ") + ":")
                _metric_table(doc, sub_rows)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))


# ---------------------------------------------------------------- CLI

async def _main(args) -> None:
    out = Path(args.out)
    out.mkdir(parents=True, exist_ok=True)
    cache = out / "jira_raw.json"

    settings = load_settings()
    http = httpx.AsyncClient(
        base_url=settings.jira_base_url,
        headers={"Authorization": f"Bearer {settings.jira_token}"},
        timeout=settings.jira_timeout_seconds, verify=settings.jira_verify_ssl,
    )
    try:
        bvs_field = args.field_id or await _discover_field_id(http, settings.jira_api_version, args.field_name)
    finally:
        await http.aclose()
    if not bvs_field:
        raise SystemExit(f"could not find '{args.field_name}' field; pass --field-id customfield_#####")
    print(f"Business Value Stream field id: {bvs_field}")

    ids = load_ticket_ids(args.tickets)
    have = json.loads(cache.read_text(encoding="utf-8")) if cache.exists() else []
    remaining = [t for t in ids if t not in {r["ticketId"] for r in have}]
    if args.use_cache and not remaining:
        records = have
        print(f"loaded {len(records)} cached tickets from {cache}")
    else:
        print(f"collecting {len(remaining)}/{len(ids)} tickets (concurrency={args.concurrency})")
        records = await collect(
            ids, bvs_field=bvs_field, concurrency=args.concurrency, ticket_timeout=args.ticket_timeout,
            done_records=have, checkpoint_path=cache, checkpoint_every=args.checkpoint_every,
        )

    stats, charts = analyze(records, out / "charts")
    (out / "stats.json").write_text(json.dumps(stats, indent=2), encoding="utf-8")
    print(f"\nstats -> {out/'stats.json'}\ncharts -> {out/'charts'}/")
    print(json.dumps(stats["usefulness"], indent=2))
    if args.docx:
        build_docx(stats, charts, out / "jira_eda.docx", n_tickets=len(records))
        print(f"docx -> {out/'jira_eda.docx'}")


if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("tickets", help="text file with one IDMT id per line")
    parser.add_argument("--out", default="out/eda/jira")
    parser.add_argument("--field-name", default="Business Value Stream")
    parser.add_argument("--field-id", default="", help="explicit customfield_##### (skips discovery)")
    parser.add_argument("--concurrency", type=int, default=6)
    parser.add_argument("--ticket-timeout", type=float, default=60.0)
    parser.add_argument("--checkpoint-every", type=int, default=20)
    parser.add_argument("--use-cache", action="store_true")
    parser.add_argument("--docx", action="store_true")
    asyncio.run(_main(parser.parse_args()))
