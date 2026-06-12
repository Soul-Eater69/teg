"""Authored Value Stream selection report — one command, no flags.

    uv run python scripts/build_vs_report.py        (needs: uv sync --extra eda --extra extract)

Builds the whole multi-chapter docx (modes, prompt engineering, historic-K) into
out/eval/vs_report.docx. The chapter structure, the written analysis, and the verdicts are
authored here; the numbers are read live from each run's <name>.runs.json so they stay current.
Charts and tables are generated per chapter. This replaces the old flag-driven comparison CLI:
to change the report, edit the CHAPTERS spec and the prose below, not a command line.
"""

from __future__ import annotations

import json
from pathlib import Path

OUT = Path("out/eval/vs_report.docx")
RUNS = Path("out/eval")


# --------------------------------------------------------------------------- data

def _mean(runs, key):
    vals = [r.get(key) for r in runs if r.get(key) is not None]
    return sum(vals) / len(vals) if vals else None


def _nested(runs, outer, inner):
    vals = [(r.get(outer) or {}).get(inner) for r in runs]
    vals = [v for v in vals if v is not None]
    return sum(vals) / len(vals) if vals else None


def _cohort(runs, prefix):
    rows, ns = [], []
    for r in runs:
        for label, prf in (r.get("cohorts") or {}).items():
            if label.strip().startswith(prefix) and prf:
                rows.append(prf)
                n = (r.get("cohort_n") or {}).get(label)
                if n is not None:
                    ns.append(n)
    if not rows:
        return None, None, None, None
    return (sum(c[0] for c in rows) / len(rows), sum(c[1] for c in rows) / len(rows),
            sum(c[2] for c in rows) / len(rows), round(sum(ns) / len(ns)) if ns else None)


def _lat_clean(runs):
    """Average latency with the single slowest ticket removed when it's a >3x-median outlier."""
    vals = []
    for r in runs:
        lat = r.get("latency") or {}
        avg, mx, med, n = lat.get("avg"), lat.get("max"), lat.get("median"), r.get("n")
        if avg is None:
            continue
        vals.append((avg * n - mx) / (n - 1) if (mx and med and n and n > 1 and mx > 3 * med) else avg)
    return sum(vals) / len(vals) if vals else None


def load(label, file, desc):
    runs = json.loads((RUNS / file).read_text(encoding="utf-8"))
    sp, sr, sf, sn = _cohort(runs, "single-VS")
    mp, mr, mf, mn = _cohort(runs, "multi-VS")
    backed, notbacked = _nested(runs, "boost", "backed_recall"), _nested(runs, "boost", "notbacked_recall")
    return {
        "label": label, "desc": desc,
        "p": _mean(runs, "micro_p"), "r": _mean(runs, "micro_r"), "f1": _mean(runs, "micro_f1"),
        "single_r": sr, "single_n": sn, "multi_r": mr, "multi_n": mn,
        "vs_r3": _nested(runs, "retrieval", "vs_lane@3"), "vs_r5": _nested(runs, "retrieval", "vs_lane@5"),
        "vs_r10": _nested(runs, "retrieval", "vs_lane@10"), "hist_r": _nested(runs, "retrieval", "historic_lane"),
        "pool_r": _nested(runs, "retrieval", "pool"),
        "backed": backed, "notbacked": notbacked,
        "lift": _nested(runs, "boost", "lift") or ((backed - notbacked) if backed and notbacked else None),
        "lat_med": _nested(runs, "latency", "median"), "lat_clean": _lat_clean(runs),
        "lat_avg": _nested(runs, "latency", "avg"), "lat_max": _nested(runs, "latency", "max"),
        "not_retrieved": _nested(runs, "buckets", "not_retrieved"),
        "gated": _nested(runs, "buckets", "gated_pre_llm"), "dropped": _nested(runs, "buckets", "llm_dropped"),
    }


# --------------------------------------------------------------------------- the report spec (authored)

CHAPTERS = [
    {
        "title": "Chapter 1 — Modes: history vs no-history, scores vs no-scores",
        "runs": [
            ("No history + scores", "all50.runs.json",
             "All 50 streams offered with the search relevance scores; no past tickets shown."),
            ("No history", "all50_noscore.runs.json",
             "All 50 streams, scores hidden; no past tickets shown."),
            ("History + scores", "evidence.runs.json",
             "All 50 streams + similar past tickets shown as examples; scores shown."),
            ("History", "evidence_noscore.runs.json",
             "All 50 streams + similar past tickets shown as examples; scores hidden."),
        ],
        "history": {"History + scores", "History"},
        "verdict": [
            "Showing the model similar PAST tickets is the dominant lever. Both history approaches beat both "
            "no-history approaches on recall by a wide margin (~0.59 → ~0.74), and they lift the hard "
            "multi-answer tickets too — so it is a real gain, not just the easy cases.",
            "The search relevance SCORES barely matter: flipping them on/off moves recall by under 0.02 either "
            "way (within run-to-run noise). The semantic ranking is weak (top-10 holds only ~26% of the "
            "correct answers), so the score is a near-useless hint. We keep the model judging on business fit "
            "plus the precedent, not the score.",
            "Verdict: ship History. Scores are optional (a wash); the past-ticket precedent is what wins.",
        ],
        "narrative": None,
    },
    {
        "title": "Chapter 2 — Prompt engineering (History, no scores)",
        "runs": [
            ("Current prompt", "evidence_noscore.runs.json",
             "The proven evidence prompt — strict 'verify each pick, don't over-trust precedent'."),
            ("Trust prompt", "history_trust.runs.json",
             "Clear fits first, then fill the requested count with precedent-backed streams instead of padding."),
            ("Recall prompt", "history_recall.runs.json",
             "Precedent as a primary signal + push completeness on multi-workflow ideas."),
        ],
        "history": {"Current prompt", "Trust prompt", "Recall prompt"},
        "verdict": [
            "The Recall prompt wins: recall 0.726 → 0.776 and hard-ticket recall 0.692 → 0.744, while the "
            "judge-precision guardrail held (0.478 → 0.470) — so the extra picks are genuinely relevant, not "
            "padding. It also captured more precedent (backed 0.76 → 0.83) and relied on it more (lift "
            "0.31 → 0.41).",
            "Trust helped too but less on hard tickets, dropped judge precision more, and ran slower. The "
            "detailed change-log below explains exactly which wording moved which number.",
        ],
        "narrative": "docs/vs_prompt_change_notes.md",
    },
    {
        "title": "Chapter 3 — Historic-K: how many past tickets to show",
        "runs": [
            ("Recall K=6", "history_recall.runs.json", "Recall prompt, 6 similar past tickets shown as evidence."),
            ("Recall K=8", "recall_k8.runs.json", "Recall prompt, 8 similar past tickets."),
            ("Recall K=10", "recall_k10.runs.json", "Recall prompt, 10 similar past tickets."),
        ],
        "history": {"Recall K=6", "Recall K=8", "Recall K=10"},
        "verdict": [
            "6 → 8 gives a small real gain (recall +0.02, hard recall +0.025). 8 → 10 is dilution, not gain: "
            "recall and hard-recall are flat, but the precedent ceiling rises while the model captures LESS of "
            "it (backed 0.838 → 0.827), lift falls (0.401 → 0.366), and judge precision keeps eroding.",
            "Decision: keep historic-K = 6. The +0.02 at K=8 is within noise and not worth the added prompt "
            "length and the dilution that grows at K=10.",
        ],
        "narrative": "docs/vs_historic_k_notes.md",
    },
]


# --------------------------------------------------------------------------- rendering

def _pct(v):
    return "n/a" if v is None else f"{v:.0%}"


def _f2(v):
    return "n/a" if v is None else f"{v:.2f}"


def _secs(v):
    return "n/a" if v is None else f"{v:.1f}s"


def _titles(ax, title, note):
    ax.figure.suptitle(title, fontsize=13, fontweight="bold", y=0.99)
    if note:
        ax.set_title(note, fontsize=9, style="italic", color="#555", pad=8)


def _bar(ax, labels, series, title, ylabel, note=None, pct=False):
    import numpy as np
    x = np.arange(len(labels))
    width = 0.8 / max(1, len(series))
    for i, (name, raw) in enumerate(series.items()):
        vals = [float("nan") if v is None else v for v in raw]
        bars = ax.bar(x + i * width - 0.4 + width / 2, vals, width, label=name)
        for b, v in zip(bars, vals):
            if v == v:
                ax.text(b.get_x() + b.get_width() / 2, v, f"{v:.0%}" if pct else f"{v:.2f}",
                        ha="center", va="bottom", fontsize=8)
    top = max((v for raw in series.values() for v in raw if v is not None), default=1.0)
    ax.set_xticks(x); ax.set_xticklabels(labels, fontsize=9); ax.set_ylabel(ylabel)
    ax.set_ylim(0, top * 1.18); _titles(ax, title, note)
    ax.legend(fontsize=8, loc="best"); ax.grid(axis="y", alpha=0.3)


def _stacked(ax, labels, series, title, ylabel, note=None):
    import numpy as np
    x = np.arange(len(labels)); bottom = np.zeros(len(labels))
    for name, raw in series.items():
        vals = np.array([0.0 if v is None else v for v in raw], dtype=float)
        bars = ax.bar(x, vals, 0.55, bottom=bottom, label=name)
        for b, v, bot in zip(bars, vals, bottom):
            if v > 0:
                ax.text(b.get_x() + b.get_width() / 2, bot + v / 2, f"{v:.0f}", ha="center",
                        va="center", fontsize=8, color="white", fontweight="bold")
        bottom += vals
    ax.set_xticks(x); ax.set_xticklabels(labels, fontsize=9); ax.set_ylabel(ylabel)
    _titles(ax, title, note); ax.legend(fontsize=8, loc="best"); ax.grid(axis="y", alpha=0.3)


def _add_table(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = "Light Grid Accent 1"
    for c, h in zip(t.rows[0].cells, headers):
        c.text = h
    for r in rows:
        cells = t.add_row().cells
        for c, v in zip(cells, r):
            c.text = str(v)


def _md(doc, text):
    """Render a small markdown subset (## / ### headings, - bullets, **bold**) into the doc."""
    def runs(p, s):
        for i, part in enumerate(s.split("**")):
            if part:
                p.add_run(part).bold = i % 2 == 1
    for line in text.splitlines():
        line = line.rstrip()
        if not line.strip():
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.lstrip().startswith(("- ", "* ")):
            runs(doc.add_paragraph(style="List Bullet"), line.lstrip()[2:])
        else:
            runs(doc.add_paragraph(), line)


def render_chapter(doc, plt, charts_dir, chapter, first):
    from docx.shared import Inches
    data = [load(*r) for r in chapter["runs"]]
    labels = [d["label"] for d in data]
    hist = chapter["history"]
    n = 0

    def save(fig):
        nonlocal n
        n += 1
        p = charts_dir / f"c{id(chapter) % 9999}_{n}.png"
        fig.savefig(p, dpi=150, bbox_inches="tight"); plt.close(fig)
        return str(p)

    if not first:
        doc.add_page_break()
    doc.add_heading(chapter["title"], level=1 if not first else 0)

    doc.add_heading("What each approach means", level=2)
    _add_table(doc, ["Approach", "What we showed the model"], [[d["label"], d["desc"]] for d in data])

    doc.add_heading("Verdict", level=2)
    for line in chapter["verdict"]:
        doc.add_paragraph(line, style="List Bullet")

    if chapter["narrative"]:
        _md(doc, Path(chapter["narrative"]).read_text(encoding="utf-8"))

    # 1. quality
    fig, ax = plt.subplots(figsize=(8, 4.6))
    _bar(ax, labels, {"Precision": [d["p"] for d in data], "Recall": [d["r"] for d in data],
                      "F1": [d["f1"] for d in data]}, "How good are the picks overall?", "fraction (0-1)",
         note="Read RECALL. Precision is low for all (10 guesses for ~2.5 real answers, by design).", pct=True)
    doc.add_heading("Precision, Recall, F1", level=2); doc.add_picture(save(fig), width=Inches(6.2))

    # 2. cohorts
    fig, ax = plt.subplots(figsize=(8, 4.6))
    _bar(ax, labels, {"Easy (1 correct answer)": [d["single_r"] for d in data],
                      "Hard (2+ correct answers)": [d["multi_r"] for d in data]},
         "Does it handle hard tickets too?", "recall", pct=True,
         note="Same metric (recall) for both. Watch the orange (hard) bar.")
    doc.add_heading("Easy vs hard tickets", level=2); doc.add_picture(save(fig), width=Inches(6.2))
    _add_table(doc, ["Approach", "Easy n", "Easy R", "Hard n", "Hard R"],
               [[d["label"], d["single_n"] or "-", _pct(d["single_r"]), d["multi_n"] or "-", _pct(d["multi_r"])]
                for d in data])

    # 3. retrieval (identical across the chapter's runs - show once)
    rep = data[0]
    fig, ax = plt.subplots(figsize=(8, 4.6))
    _bar(ax, ["Search\ntop 3", "Search\ntop 5", "Search\ntop 10", "Past-ticket\nexamples", "Full catalogue"],
         {"correct answers present": [rep["vs_r3"], rep["vs_r5"], rep["vs_r10"], rep["hist_r"], rep["pool_r"]]},
         "Did the system put the right answers in front of the model?", "recall", pct=True,
         note="The full catalogue is always shown, so the ceiling is ~100%. Search ranking alone is weak.")
    doc.add_heading("Were the right answers shown to the model?", level=2)
    doc.add_picture(save(fig), width=Inches(6.2))

    # 4. where misses go
    fig, ax = plt.subplots(figsize=(8, 4.6))
    _stacked(ax, labels, {"Never retrieved (system)": [d["not_retrieved"] for d in data],
                          "Filtered before model (system)": [d["gated"] for d in data],
                          "Model saw it, didn't pick it": [d["dropped"] for d in data]},
             "Where did the MISSED answers go?", "missed answers",
             note="All in the top band = retrieval found everything; only the prompt can improve it.")
    doc.add_picture(save(fig), width=Inches(6.2))

    # 5. precedent (history runs only)
    hd = [d for d in data if d["label"] in hist and d["backed"] is not None]
    if hd:
        fig, ax = plt.subplots(figsize=(8, 4.6))
        _bar(ax, [d["label"] for d in hd],
             {"Answers that WERE in the past tickets": [d["backed"] for d in hd],
              "Answers that were NOT in them": [d["notbacked"] for d in hd]},
             "When we show past tickets, does the model follow them?", "how often the model picked it",
             note="Tall blue + short orange = the model leans on precedent. The gap is 'lift'.", pct=True)
        doc.add_heading("Precedent: backed vs lift", level=2); doc.add_picture(save(fig), width=Inches(6.2))
        _add_table(doc, ["Approach", "Precedent backed", "Precedent lift"],
                   [[d["label"], _pct(d["backed"]), f"{d['lift']:+.2f}" if d["lift"] is not None else "n/a"]
                    for d in hd])

    # 6. latency
    fig, ax = plt.subplots(figsize=(8, 4.6))
    _bar(ax, labels, {"typical (median)": [d["lat_med"] for d in data],
                      "average (outlier removed)": [d["lat_clean"] for d in data]},
         "How long does one prediction take?", "seconds",
         note="Median = typical speed. The average has the one retry-storm ticket removed so the scale is real.")
    doc.add_heading("Latency", level=2); doc.add_picture(save(fig), width=Inches(6.2))

    # summary table
    _add_table(doc, ["Approach", "P", "R", "F1", "Easy R", "Hard R", "Backed", "Lift", "Typical"],
               [[d["label"], _pct(d["p"]), _pct(d["r"]), _f2(d["f1"]), _pct(d["single_r"]), _pct(d["multi_r"]),
                 _pct(d["backed"]), f"{d['lift']:+.2f}" if d["lift"] is not None else "n/a", _secs(d["lat_med"])]
                for d in data])


GLOSSARY = [
    ("Recall", "Of the correct Value Streams, the fraction we found. Higher is better."),
    ("Precision", "Of our guesses, the fraction correct. Low by design (10 guesses, ~2.5 real)."),
    ("F1", "Balance of precision and recall. Higher is better."),
    ("Easy / Hard ticket", "Easy = one correct stream; Hard = two or more. Tracked apart so a win isn't "
                           "just the easy cases."),
    ("Precedent backed", "Of the correct answers that were in the shown past tickets, the fraction the model "
                         "picked. How much of the offered precedent it used."),
    ("Precedent lift", "(pick-rate on answers IN the past tickets) - (pick-rate on answers NOT in them). "
                       "Whether the precedent is what's causing the pick."),
    ("Ceiling", "Fraction of correct answers on the candidate list at all (~100%, the full catalogue is shown)."),
    ("Where misses go", "Never retrieved / filtered before the model / model saw it and dropped it. The last "
                        "is the model's choice, not the system's."),
]


def main():
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    from docx import Document

    charts_dir = OUT.parent / "report_charts"
    charts_dir.mkdir(parents=True, exist_ok=True)
    doc = Document()
    doc.add_heading("Value Stream selection — tuning report", level=0)
    doc.add_paragraph(
        "How we tune the model that predicts a ticket's Value Streams. Three chapters, each on the same "
        "60-ticket eval set: (1) which mode wins, (2) prompt engineering on that mode, (3) how many past "
        "tickets to show. Every chart states its question and how to read it; a metric glossary is at the end.")
    for i, chapter in enumerate(CHAPTERS):
        render_chapter(doc, plt, charts_dir, chapter, first=(i == 0))

    doc.add_page_break()
    doc.add_heading("What the metrics mean", level=1)
    _add_table(doc, ["Term", "Plain meaning"], [[t, m] for t, m in GLOSSARY])

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print(f"report -> {OUT}")


if __name__ == "__main__":
    main()
