"""Batch evaluation of Value Stream prediction against the historic ground truth.

Dataset = the ingested Cosmos IDMT docs (out/idmt/cosmos_idmt.json): each carries the
condensed summaryFields (input) and its approved VS ground truth (properties.themes[]).

Per ticket: run prediction, EXCLUDING the ticket itself from the historic analog lane
(leave-one-out - a ticket must not see its own GT). Compare the predicted VS ids against
the GT ids. Reports precision / recall / F1 (micro + macro) and precision@k / recall@k.

Experiments:
  (default)            condensed summaryFields, with direct/implied classification
  --raw-text           feed properties.rawText as the query instead of summaryFields

Usage:
  uv run python -m scripts.eval_vs out/idmt/cosmos_idmt.json --count 10 --k 3 5 10
"""

from __future__ import annotations

import argparse
import asyncio
import csv
import json
from pathlib import Path
from time import perf_counter

from teg.bootstrap import build_value_stream_service
from teg.config.settings import load_settings
from teg.contracts.value_stream_io import ValueStreamRequest
from teg.domain.condensed import SummaryFields
from teg.integrations.llm import build_llm_client
from teg.integrations.search import HistoricalValueStreamLabel
from teg.value_stream.config import ValueStreamConfig
from teg.value_stream.drop_explainer import explain_drops, explain_swaps, score_candidates
from teg.value_stream.relevance_judge import judge_value_streams


def _load(path: str) -> list[dict]:
    docs = json.loads(Path(path).read_text(encoding="utf-8"))
    docs = docs if isinstance(docs, list) else [docs]
    _attach_gt_from_themes(docs, Path(path))
    return docs


def _attach_gt_from_themes(idmt_docs: list[dict], idmt_path: Path) -> None:
    """Reconstruct each ER's properties.themes (the VS ground truth) from the sibling themes file.

    The Cosmos schema drops properties.themes from the ER doc - the GT lives in the separate Theme
    docs (one per linked theme, parentRef == the ER's sourceId, properties.valueStream the tagged VS).
    Rebuild it here so the rest of the eval (which reads properties.themes) is unchanged. No-op when
    the docs already carry themes (old-format file) or the themes file is absent.
    """
    if any((d.get("properties") or {}).get("themes") for d in idmt_docs):
        return  # GT already embedded
    themes_path = idmt_path.with_name("cosmos_themes.json")
    if not themes_path.exists():
        return
    theme_docs = json.loads(themes_path.read_text(encoding="utf-8"))
    by_parent: dict[str, list[dict]] = {}
    for t in theme_docs:
        parent = t.get("parentRef")
        vs = (t.get("properties") or {}).get("valueStream") or {}
        if parent and vs.get("valueStreamId"):
            by_parent.setdefault(parent, []).append({
                "key": t.get("key"),
                "sourceId": t.get("sourceId"),
                "valueStreamId": vs.get("valueStreamId"),
                "valueStreamName": vs.get("valueStreamName"),
            })
    for doc in idmt_docs:
        doc.setdefault("properties", {})["themes"] = by_parent.get(doc.get("sourceId"), [])


_EMBED_TOKEN_CAP = 8000  # under the embedding model's ~8191-token hard limit
_ENC = None


def _truncate_tokens(text: str, max_tokens: int) -> str:
    """Truncate to actual cl100k_base tokens (the text-embedding-3 tokenizer); lazy-loaded."""
    global _ENC
    if _ENC is None:
        import tiktoken
        _ENC = tiktoken.get_encoding("cl100k_base")
    toks = _ENC.encode(text)
    return text if len(toks) <= max_tokens else _ENC.decode(toks[:max_tokens])


def _summary_fields(props: dict, *, raw_text: bool, query_budget: int = 0) -> SummaryFields:
    # New schema: the LLM summary is businessSummary (properties.summary is the ticket title).
    # Fall back to summary for any pre-rename file.
    llm_summary = props.get("businessSummary") or props.get("summary", "")
    if raw_text:
        text = props.get("rawText", "") or llm_summary
        # This text is EMBEDDED for retrieval (raw_text here = the retrieval representation), so it
        # must stay under the model's hard ~8191-token cap. Token-accurate truncate (a chars
        # heuristic overshoots on dense text and 400s the embeddings endpoint).
        text = _truncate_tokens(text, min(query_budget or _EMBED_TOKEN_CAP, _EMBED_TOKEN_CAP))
        return SummaryFields(generated_summary=text, business_problem="", business_capability="")
    return SummaryFields(
        generated_summary=llm_summary,
        business_problem=props.get("businessProblem", ""),
        business_capability=props.get("businessCapability", ""),
        key_terms=props.get("keyTerms", []) or [],
        stakeholders=props.get("stakeholders", []) or [],
        systems_and_products=props.get("systemsAndProducts", []) or [],
    )


def _prompt_text(props: dict, *, raw_text: bool, query_budget: int) -> str:
    """Raw ticket text for the SELECTION prompt only (empty -> the prompt uses the summary).

    Retrieval ALWAYS uses the summary (the embeddable matcher); this changes only what the LLM that
    picks the VS reads - so '--raw-text' gives the prompt raw context WITHOUT degrading retrieval.
    """
    if not raw_text:
        return ""
    text = props.get("rawText", "") or ""
    return text[:query_budget * 4] if query_budget else text  # ~4 chars/token


def _gt_ids(props: dict) -> set[str]:
    return {t["valueStreamId"] for t in (props.get("themes") or []) if t.get("valueStreamId")}


def _cardinality_stats(rows: list[dict]) -> dict:
    """Predicted VS count vs GT count: exact/under/over-fetch rates + the delta distribution.

    exact_count_rate = predicted_count == gt_count (right SIZE, regardless of which VS).
    exact_set_rate   = the predicted SET equals GT (fp == fn == 0; the strong version).
    under/over       = predicted fewer / more than GT. delta_dist histograms (pred - gt).
    """
    n = len(rows)
    exact_count = under = over = exact_set = 0
    deltas: list[int] = []
    for r in rows:
        d = r["predicted_count"] - r["gt_count"]
        deltas.append(d)
        exact_count += d == 0
        under += d < 0
        over += d > 0
        exact_set += r["fp"] == 0 and r["fn"] == 0
    dist: dict[str, int] = {}
    for d in deltas:
        dist[f"{d:+d}"] = dist.get(f"{d:+d}", 0) + 1
    return {
        "n": n,
        "exact_count_rate": _div(exact_count, n),
        "under_rate": _div(under, n),
        "over_rate": _div(over, n),
        "exact_set_rate": _div(exact_set, n),
        "avg_delta": _div(sum(deltas), n),
        "avg_abs_delta": _div(sum(abs(d) for d in deltas), n),
        "delta_dist": dict(sorted(dist.items(), key=lambda kv: int(kv[0]))),
    }


def _count_following_stats(rows: list[dict]) -> dict:
    """Did the LLM return the REQUESTED count on its own, before our _enforce_count padding?

    Compares the LLM's raw pick count (trace.llm_pick_count) to requested_count. 'followed' = it
    returned exactly the count; 'under' = it picked fewer (we then padded - 'padded_rate'); 'over' =
    it picked more (we trimmed). This is the honest count-adherence signal - the final
    predicted_count is forced by us, so it always matches and hides the model's behaviour.
    """
    usable = [r for r in rows if r.get("requested_count")]
    if not usable:
        return {}
    n = len(usable)
    followed = under = over = 0
    deltas: list[int] = []
    for r in usable:
        d = r["llm_pick_count"] - r["requested_count"]
        deltas.append(d)
        followed += d == 0
        under += d < 0
        over += d > 0
    dist: dict[str, int] = {}
    for d in deltas:
        dist[f"{d:+d}"] = dist.get(f"{d:+d}", 0) + 1
    return {
        "n": n,
        "followed_rate": _div(followed, n),
        "under_rate": _div(under, n),
        "over_rate": _div(over, n),
        "padded_rate": _div(under, n),  # under-picks are the ones we had to pad up
        "avg_requested": _div(sum(r["requested_count"] for r in usable), n),
        "avg_llm_pick": _div(sum(r["llm_pick_count"] for r in usable), n),
        "delta_dist": dict(sorted(dist.items(), key=lambda kv: int(kv[0]))),
    }


def _cohort_prf(rows: list[dict]) -> tuple[float, float, float]:
    """Micro P/R/F1 over a cohort of per-ticket rows (each has tp/fp/fn)."""
    tp = sum(r["tp"] for r in rows)
    fp = sum(r["fp"] for r in rows)
    fn = sum(r["fn"] for r in rows)
    p, r = _div(tp, tp + fp), _div(tp, tp + fn)
    return p, r, _div(2 * p * r, p + r)


def _load_no_attachment_ids(cache_path: str) -> set[str]:
    """Ticket ids with zero attachments, from an EDA attachments_raw.json cache (optional)."""
    try:
        records = json.loads(Path(cache_path).read_text(encoding="utf-8"))
    except Exception:
        return set()
    has_att: set[str] = set()
    seen: set[str] = set()
    for r in records:
        seen.add(r.get("ticketId", ""))
        if r.get("kind") == "attachment":
            has_att.add(r.get("ticketId", ""))
    return {t for t in seen if t and t not in has_att}


def _vs_name_map(docs: list[dict]) -> dict[str, str]:
    names: dict[str, str] = {}
    for doc in docs:
        for t in doc.get("properties", {}).get("themes") or []:
            if t.get("valueStreamId"):
                names[t["valueStreamId"]] = t.get("valueStreamName") or t["valueStreamId"]
    return names


def _base_rate_counts(docs: list[dict]) -> tuple[dict[str, int], int]:
    """Corpus tag frequency: how many tickets carry each VS as GT, and the ticket total.

    The 'breadth' prior - a VS tagged on many tickets is broad/generic. Built once over the
    whole corpus; the per-ticket rate is computed leave-one-out so a ticket never informs its
    own penalty.
    """
    counts: dict[str, int] = {}
    total = 0
    for doc in docs:
        gt = _gt_ids(doc.get("properties", {}))
        if not gt:
            continue
        total += 1
        for vs in gt:
            counts[vs] = counts.get(vs, 0) + 1
    return counts, total


def _loo_base_rates(counts: dict[str, int], total: int, gt: set[str]) -> dict[str, float]:
    """Per-ticket base rate excluding this ticket's own tags (leave-one-out)."""
    denom = max(1, total - 1)
    return {vs: (n - (1 if vs in gt else 0)) / denom for vs, n in counts.items()}


async def _collect_predictions(service, args, jobs, sem) -> dict[str, list[str]]:
    """Pass 1: predict every ticket with the penalty OFF, to measure each stream's FP-rate."""
    async def _one(doc, ticket_id, gt):
        async with sem:
            req = ValueStreamRequest(
                ticket_id=ticket_id,
                summary_fields=_summary_fields(doc.get("properties", {}), raw_text=args.raw_retrieval, query_budget=args.retrieval_budget),  # retrieval repr
                prompt_text=_prompt_text(doc.get("properties", {}), raw_text=args.raw_text, query_budget=args.query_budget),
                requested_count=_requested_count(args, gt),
                exclude_ticket_ids=[ticket_id],
            )
            try:
                resp, _ = await service.predict_traced(req, base_rates={})  # {} = no penalty
                return ticket_id, [r.value_stream_id for r in resp.recommendations]
            except Exception:
                return ticket_id, []
    print(f"FP-rate prior: pass 1 predicting {len(jobs)} tickets (penalty off)...")
    pairs = await asyncio.gather(*(_one(d, t, g) for d, t, g, _ in jobs))
    return dict(pairs)


def _fp_rate_stats(pred_by_ticket: dict[str, list[str]], gt_by_ticket: dict[str, set[str]]):
    """Per-VS (predicted_count, false_positive_count) across the corpus + the per-ticket views."""
    pred_count: dict[str, int] = {}
    fp_count: dict[str, int] = {}
    for tid, preds in pred_by_ticket.items():
        gt = gt_by_ticket.get(tid, set())
        for vs in set(preds):
            pred_count[vs] = pred_count.get(vs, 0) + 1
            if vs not in gt:
                fp_count[vs] = fp_count.get(vs, 0) + 1
    return pred_count, fp_count


def _loo_fp_rates(pred_count, fp_count, ticket_preds: list[str], gt: set[str]) -> dict[str, float]:
    """Per-ticket FP-rate prior, leave-one-out (Laplace-smoothed so rarely-predicted ~0.5).

    fp_rate = (this stream's false positives) / (times it was predicted). High = attractor.
    """
    preds = set(ticket_preds)
    rates: dict[str, float] = {}
    for vs, p in pred_count.items():
        p_loo = p - (1 if vs in preds else 0)
        fp_loo = fp_count.get(vs, 0) - (1 if (vs in preds and vs not in gt) else 0)
        rates[vs] = (fp_loo + 1) / (p_loo + 2)  # smoothed; neutral 0.5 when unseen
    return rates


def _prf(predicted: list[str], gt: set[str]) -> tuple[int, int, int]:
    pset = set(predicted)
    tp = len(pset & gt)
    return tp, len(pset) - tp, len(gt) - tp  # tp, fp, fn


def _div(a: float, b: float) -> float:
    return a / b if b else 0.0


def _requested_count(args, gt: set[str]) -> int:
    if args.count_mode == "gt":
        return max(1, len(gt))
    if args.count_mode == "gt_buffer":
        return max(1, len(gt) + args.buffer)
    return args.count


async def _eval_one(service, llm, args, doc, ticket_id: str, gt: set[str], base_rates, vs_names, sem, progress) -> dict:
    async with sem:
        request = ValueStreamRequest(
            ticket_id=ticket_id,
            summary_fields=_summary_fields(doc.get("properties", {}), raw_text=args.raw_retrieval, query_budget=args.retrieval_budget),  # retrieval repr
            prompt_text=_prompt_text(doc.get("properties", {}), raw_text=args.raw_text, query_budget=args.query_budget),
            requested_count=_requested_count(args, gt),
            exclude_ticket_ids=[ticket_id],  # leave-one-out
        )
        try:
            t0 = perf_counter()
            resp, trace = await service.predict_traced(request, base_rates=base_rates)
            elapsed = perf_counter() - t0  # prediction latency only (excludes judge/explain probes)
            predicted = [r.value_stream_id for r in resp.recommendations]
        except Exception as exc:  # one bad ticket must not abort the batch
            progress["done"] += 1
            print(f"[{progress['done']}/{progress['total']}] {ticket_id}  ERROR {type(exc).__name__}: {exc}")
            return {"ticket_id": ticket_id, "gt": gt, "predicted": [], "error": str(exc)}
        buckets = _miss_buckets(gt, predicted, trace)
        retrieval = _retrieval_recall(gt, trace, args.k)
        boost = _historic_boost(gt, predicted, trace)
        # Post-hoc: ask why the LLM dropped GT it actually saw (never changes the metrics).
        drop_reasons: dict[str, str] = {}
        if llm is not None and buckets["llm_dropped"]:
            try:
                explained = await explain_drops(
                    query=request.summary_fields.generated_summary,
                    review_pool=trace.review_pool,
                    picked_ids=predicted,
                    dropped_ids=buckets["llm_dropped"],
                    llm_client=llm,
                )
                drop_reasons = {vs: exp.reason_code for vs, exp in explained.items()}
            except Exception as exc:  # a failed probe must not abort the batch
                print(f"    explain-drops failed for {ticket_id}: {type(exc).__name__}: {exc}")
        # Level B: independent 0-1 score of every candidate -> how close each dropped GT was to the cut.
        score_margins: dict[str, float] = {}
        if llm is not None and args.score_margins and buckets["llm_dropped"]:
            try:
                scores = await score_candidates(
                    query=request.summary_fields.generated_summary,
                    review_pool=trace.review_pool, llm_client=llm,
                )
                # Cut = lowest self-score among the picks; margin = dropped GT score - cut.
                cut = min((scores.get(p, 0.0) for p in predicted), default=0.0)
                score_margins = {vs: round(scores.get(vs, 0.0) - cut, 3) for vs in buckets["llm_dropped"]}
            except Exception as exc:
                print(f"    score-margins failed for {ticket_id}: {type(exc).__name__}: {exc}")
        # Level C: comparative - why the picks beat each dropped GT (richer swap taxonomy).
        swap_reasons: dict[str, str] = {}
        if llm is not None and args.explain_swaps and buckets["llm_dropped"]:
            try:
                swaps = await explain_swaps(
                    query=request.summary_fields.generated_summary,
                    review_pool=trace.review_pool, picked_ids=predicted,
                    dropped_ids=buckets["llm_dropped"], llm_client=llm,
                )
                swap_reasons = {vs: exp.reason_code for vs, exp in swaps.items()}
            except Exception as exc:
                print(f"    explain-swaps failed for {ticket_id}: {type(exc).__name__}: {exc}")
        # LLM-as-judge: is each predicted / missed VS genuinely relevant (independent of GT)?
        judged: dict[str, bool] = {}
        if llm is not None and args.judge:
            try:
                pool_desc = {c.value_stream_id: c.value_stream_description for c in trace.review_pool}
                name_of = {**vs_names, **{r.value_stream_id: r.value_stream_name for r in resp.recommendations}}
                ids = set(predicted) | gt
                items = [(i, name_of.get(i, i), pool_desc.get(i, "")) for i in ids]
                judged = await judge_value_streams(
                    query=request.summary_fields.generated_summary, items=items, llm_client=llm
                )
            except Exception as exc:
                print(f"    judge failed for {ticket_id}: {type(exc).__name__}: {exc}")
    tp, fp, fn = _prf(predicted, gt)
    progress["done"] += 1
    print(f"[{progress['done']}/{progress['total']}] {ticket_id}  "
          f"P={_div(tp, tp+fp):.2f} R={_div(tp, tp+fn):.2f}  (gt={len(gt)}, pred={len(predicted)}, {elapsed:.1f}s)")
    return {"ticket_id": ticket_id, "gt": gt, "predicted": predicted, "buckets": buckets,
            "drop_reasons": drop_reasons, "judged": judged, "elapsed": elapsed,
            "score_margins": score_margins, "swap_reasons": swap_reasons,
            "retrieval_s": trace.retrieval_seconds, "selection_s": trace.selection_seconds,
            "llm_pick_count": trace.llm_pick_count, "requested_count": trace.requested_count,
            "retrieval": retrieval, "boost": boost}


def _historic_boost(gt: set[str], predicted: list[str], trace) -> dict:
    """Did appearing in the historic evidence make a GT value stream more likely to be picked?

    Splits this ticket's GT into 'backed by historic' (the VS appeared in the 6 similar tickets'
    evidence) vs 'not backed', and counts recall on each. Aggregated, the gap between the two recalls
    is the lift the historic evidence adds to selection (most meaningful in evidence/merge modes).
    """
    hist = set(trace.historic_lane_ids)
    pred = set(predicted)
    backed = gt & hist
    not_backed = gt - hist
    return {
        "backed_total": len(backed), "backed_hit": len(backed & pred),
        "notbacked_total": len(not_backed), "notbacked_hit": len(not_backed & pred),
    }


def _retrieval_recall(gt: set[str], trace, ks: list[int]) -> dict:
    """Per-lane retrieval recall (the ceiling, independent of the LLM):
      vs_lane@k     - GT in the top-k of the semantic VS ranking
      historic_lane - GT surfaced by the historic lane's candidates
      pool          - GT that made it into the review pool the LLM saw
    """
    n = max(1, len(gt))
    vs_ranked = trace.vs_lane_ranked
    out = {f"vs_lane@{k}": len(gt & set(vs_ranked[:k])) / n for k in ks}
    out["historic_lane"] = len(gt & set(trace.historic_lane_ids)) / n
    out["pool"] = len(gt & set(trace.review_pool_ids)) / n
    return out


def _miss_buckets(gt: set[str], predicted: list[str], trace) -> dict[str, list[str]]:
    """Localize each FN: where did the right answer die?

    not_retrieved  -> never made the merged candidate set (retrieval gap)
    gated_pre_llm  -> retrieved, but the merger dropped it before the LLM saw it
    llm_dropped    -> the LLM saw it in the review pool and still didn't pick it
    """
    retrieved = set(trace.retrieved_ids)
    pool = set(trace.review_pool_ids)
    out: dict[str, list[str]] = {"not_retrieved": [], "gated_pre_llm": [], "llm_dropped": []}
    for vs in gt - set(predicted):
        if vs not in retrieved:
            out["not_retrieved"].append(vs)
        elif vs not in pool:
            out["gated_pre_llm"].append(vs)
        else:
            out["llm_dropped"].append(vs)
    return out


async def main(args) -> None:
    docs = _load(args.dataset)
    # Mode -> window default (all50/evidence see the full catalogue; topk uses --window or 18).
    window = args.window or {"all50": 50, "evidence": 50}.get(args.mode, 0)
    config = ValueStreamConfig(
        use_historic_lane=not args.semantic_only,
        generic_penalty_scale=args.generic_penalty,
        min_confidence=args.min_confidence,
        selection_mode=args.mode,
        selection_prompt_override=args.selection_prompt,
        show_candidate_scores=not args.no_candidate_scores,
        historic_repr=args.historic_repr,
        historic_budget=args.historic_budget,
        **({"llm_candidate_window": window} if window else {}),
        **({"historical_fetch_k": args.historic_k} if args.historic_k else {}),
    )
    # Local content lookup so the evidence block + VS labels come from the corpus (keyed by business
    # key) - the offline stand-in for production's Cosmos point-reads (the index is retrieval-only).
    historic_content = {
        d.get("key", ""): {"raw": (d.get("properties", {}).get("rawText") or ""),
                           "description": (d.get("properties", {}).get("description") or ""),
                           "summary": (d.get("properties", {}).get("businessSummary") or ""),
                           "vs": [HistoricalValueStreamLabel(t["valueStreamId"], t.get("valueStreamName", ""))
                                  for t in (d.get("properties", {}).get("themes") or [])
                                  if t.get("valueStreamId")]}
        for d in docs
    }
    service = build_value_stream_service(config=config, historic_content=historic_content)
    llm = build_llm_client(load_settings()) if (
        args.explain_drops or args.judge or args.score_margins or args.explain_swaps) else None
    sem = asyncio.Semaphore(args.concurrency)

    vs_names = _vs_name_map(docs)  # id -> name, for judging GT streams not in the pool
    base_jobs = []
    skipped = 0
    for i, doc in enumerate(docs, start=1):
        gt = _gt_ids(doc.get("properties", {}))
        if len(gt) < args.min_gt:  # drop tickets with too few GT VS (e.g. single-label)
            skipped += 1
            continue
        base_jobs.append((doc, doc.get("key") or doc.get("ticketId") or doc.get("sourceId") or f"row{i}", gt))

    # Sample a fixed subset for fast prompt iteration. Seeded + sorted by ticket id so every run
    # (strong vs lean, summary vs raw, repeat 1/2/3) hits the SAME tickets -> comparable.
    if args.sample and args.sample < len(base_jobs):
        import random
        rng = random.Random(args.seed)
        base_jobs = sorted(rng.sample(base_jobs, args.sample), key=lambda j: j[1])
        print(f"sampled {len(base_jobs)} tickets (seed={args.seed}) for this run")

    # Build the per-ticket penalty prior (leave-one-out) from the chosen signal.
    #   gt_freq = corpus tag frequency (broad = often a GT). fp_rate = false-positive rate
    #   (broad = often predicted-but-wrong; the correct attractor signal, needs a pass 1).
    rate_for: dict[str, dict[str, float] | None] = {t: None for _, t, _ in base_jobs}
    if args.generic_penalty > 0:
        if args.penalty_signal == "fp_rate":
            gt_by = {t: g for _, t, g in base_jobs}
            preds = await _collect_predictions(service, args, [(d, t, g, None) for d, t, g in base_jobs], sem)
            pred_count, fp_count = _fp_rate_stats(preds, gt_by)
            for _, t, g in base_jobs:
                rate_for[t] = _loo_fp_rates(pred_count, fp_count, preds.get(t, []), g)
        else:
            counts, total = _base_rate_counts(docs)
            for _, t, g in base_jobs:
                rate_for[t] = _loo_base_rates(counts, total, g)

    jobs = [(d, t, g, rate_for[t]) for d, t, g in base_jobs]
    progress = {"done": 0, "total": len(jobs)}
    print(f"evaluating {len(jobs)} tickets (concurrency={args.concurrency}; "
          f"skipped {skipped} with < {args.min_gt} GT value streams; "
          f"penalty={args.generic_penalty} signal={args.penalty_signal})")
    try:
        results = await asyncio.gather(
            *(_eval_one(service, llm, args, d, t, g, r, vs_names, sem, progress) for d, t, g, r in jobs)
        )
    finally:
        await service.aclose()
        if llm is not None and hasattr(llm, "aclose"):
            await llm.aclose()

    rows: list[dict] = []
    micro_tp = micro_fp = micro_fn = 0
    bucket_totals = {"not_retrieved": 0, "gated_pre_llm": 0, "llm_dropped": 0}
    reason_totals: dict[str, int] = {}
    swap_totals: dict[str, int] = {}  # Level C reason codes
    margins: list[float] = []  # Level B: dropped-GT score minus the pick cut
    # Judge-adjusted: predictions judged relevant (even if not GT) + misses judged supported.
    judge_pred = judge_rel_pred = judge_supported_miss = 0
    retrieval_acc: dict[str, list] = {}
    boost_acc = {"backed_total": 0, "backed_hit": 0, "notbacked_total": 0, "notbacked_hit": 0}
    p_at = {k: [] for k in args.k}
    r_at = {k: [] for k in args.k}
    for res in results:
        if res.get("error"):
            continue
        for key, val in (res.get("retrieval") or {}).items():
            retrieval_acc.setdefault(key, []).append(val)
        for key in boost_acc:
            boost_acc[key] += (res.get("boost") or {}).get(key, 0)
        gt, predicted = res["gt"], res["predicted"]
        tp, fp, fn = _prf(predicted, gt)
        micro_tp += tp; micro_fp += fp; micro_fn += fn
        judged = res.get("judged") or {}
        if judged:
            judge_pred += len(predicted)
            judge_rel_pred += sum(1 for p in predicted if judged.get(p))
            judge_supported_miss += sum(1 for g in gt if g not in predicted and judged.get(g))
        buckets = res.get("buckets") or {}
        for name in bucket_totals:
            bucket_totals[name] += len(buckets.get(name, []))
        drop_reasons = res.get("drop_reasons") or {}
        for code in drop_reasons.values():
            reason_totals[code] = reason_totals.get(code, 0) + 1
        for code in (res.get("swap_reasons") or {}).values():
            swap_totals[code] = swap_totals.get(code, 0) + 1
        margins.extend((res.get("score_margins") or {}).values())
        for k in args.k:
            topk = set(predicted[:k])
            p_at[k].append(_div(len(topk & gt), min(k, len(predicted)) or 1))
            r_at[k].append(_div(len(topk & gt), len(gt)))
        rows.append({
            "ticket_id": res["ticket_id"], "gt_count": len(gt), "predicted_count": len(predicted),
            "tp": tp, "fp": fp, "fn": fn, "seconds": round(res.get("elapsed", 0.0), 2),
            "retrieval_s": round(res.get("retrieval_s", 0.0), 2),
            "selection_s": round(res.get("selection_s", 0.0), 2),
            "llm_pick_count": res.get("llm_pick_count", 0),
            "requested_count": res.get("requested_count", 0),
            "precision": round(_div(tp, tp + fp), 3), "recall": round(_div(tp, tp + fn), 3),
            "fn_not_retrieved": "; ".join(buckets.get("not_retrieved", [])),
            "fn_gated_pre_llm": "; ".join(buckets.get("gated_pre_llm", [])),
            "fn_llm_dropped": "; ".join(buckets.get("llm_dropped", [])),
            "fn_drop_reasons": "; ".join(f"{vs}={code}" for vs, code in drop_reasons.items()),
            "fn_swap_reasons": "; ".join(f"{vs}={code}" for vs, code in (res.get("swap_reasons") or {}).items()),
            "fn_score_margins": "; ".join(f"{vs}={m:+.2f}" for vs, m in (res.get("score_margins") or {}).items()),
            "gt": "; ".join(sorted(gt)), "predicted": "; ".join(predicted),
        })

    n = len(rows)
    micro_p = _div(micro_tp, micro_tp + micro_fp)
    micro_r = _div(micro_tp, micro_tp + micro_fn)
    macro_p = _div(sum(r["precision"] for r in rows), n)
    macro_r = _div(sum(r["recall"] for r in rows), n)

    print("\n" + "=" * 60)
    print(f"tickets evaluated: {n}   "
          f"(mode={args.mode}, count_mode={args.count_mode}, "
          f"input={'rawText' if args.raw_text else 'condensed'}, window={args.window or 18}, "
          f"generic_penalty={args.generic_penalty}/{args.penalty_signal if args.generic_penalty else '-'})")
    avg_pred = _div(sum(r["predicted_count"] for r in rows), n)
    avg_gt = _div(sum(r["gt_count"] for r in rows), n)
    print(f"avg predicted={avg_pred:.1f}  avg gt={avg_gt:.1f}  "
          f"(min_confidence={args.min_confidence}{' = abstention on' if args.min_confidence else ''})")
    cardinality = _cardinality_stats(rows)
    print(f"cardinality (predicted count vs gt count; count_mode={args.count_mode}):")
    print(f"  exact={cardinality['exact_count_rate']:.0%}  under={cardinality['under_rate']:.0%}  "
          f"over={cardinality['over_rate']:.0%}  | exact set={cardinality['exact_set_rate']:.0%}  "
          f"avg delta={cardinality['avg_delta']:+.2f}")
    print(f"  delta dist (pred-gt): {cardinality['delta_dist']}")
    follow = _count_following_stats(rows)
    if follow:
        print(f"count-following (did the LLM return what was requested, BEFORE our padding):")
        print(f"  followed={follow['followed_rate']:.0%}  under={follow['under_rate']:.0%}  "
              f"over={follow['over_rate']:.0%}  | avg requested={follow['avg_requested']:.1f}  "
              f"avg LLM picked={follow['avg_llm_pick']:.1f}  padded={follow['padded_rate']:.0%}")
        print(f"  LLM(picked-requested) dist: {follow['delta_dist']}")
    print(f"micro  P={micro_p:.3f}  R={micro_r:.3f}  F1={_div(2*micro_p*micro_r, micro_p+micro_r):.3f}  (strict GT)")
    print(f"macro  P={macro_p:.3f}  R={macro_r:.3f}  F1={_div(2*macro_p*macro_r, macro_p+macro_r):.3f}  (strict GT)")
    for k in args.k:
        print(f"  @{k:<2}  P@{k}={_div(sum(p_at[k]), n):.3f}   R@{k}={_div(sum(r_at[k]), n):.3f}")

    retrieval_mean = {k: _div(sum(v), len(v)) for k, v in retrieval_acc.items()}
    if retrieval_mean:
        print("\nretrieval recall (the ceiling - did retrieval put GT in front of the LLM?):")
        for k in args.k:
            key = f"vs_lane@{k}"
            if key in retrieval_mean:
                print(f"  VS lane R@{k:<2} = {retrieval_mean[key]:.3f}  (GT in top-{k} of semantic ranking)")
        used = "used in pool" if args.mode in ("merge", "historic_only", "evidence") else "DIAGNOSTIC - not used in this mode"
        print(f"  historic lane R = {retrieval_mean.get('historic_lane', 0):.3f}  (GT surfaced by historic precedent; {used})")
        print(f"  review pool R   = {retrieval_mean.get('pool', 0):.3f}  (GT the LLM actually saw - the ceiling)")

    # Historic boost: recall on GT that the historic evidence surfaced vs GT it didn't. The gap is
    # how much the historic signal lifts selection (most meaningful in evidence/merge).
    br = _div(boost_acc["backed_hit"], boost_acc["backed_total"])
    nr_ = _div(boost_acc["notbacked_hit"], boost_acc["notbacked_total"])
    print("\nhistoric boost (recall split by whether the GT appeared in the historic evidence):")
    print(f"  GT backed by historic:     {boost_acc['backed_hit']}/{boost_acc['backed_total']} = {br:.3f} recall")
    print(f"  GT NOT in historic:        {boost_acc['notbacked_hit']}/{boost_acc['notbacked_total']} = {nr_:.3f} recall")
    print(f"  -> lift from historic = {br - nr_:+.3f}  (higher recall when the GT is in the evidence)")

    # Cohort breakdown: single-VS (gt==1, the easy half excluded by --min-gt 2) vs multi-VS.
    no_att = _load_no_attachment_ids(args.attachments_cache) if args.attachments_cache else set()
    cohorts = [
        ("single-VS (gt==1)", [r for r in rows if r["gt_count"] == 1]),
        ("multi-VS  (gt>=2)", [r for r in rows if r["gt_count"] >= 2]),
    ]
    if no_att:
        cohorts.append(("no-attachment   ", [r for r in rows if r["ticket_id"] in no_att]))
        cohorts.append(("no-att + single ", [r for r in rows if r["ticket_id"] in no_att and r["gt_count"] == 1]))
    print("\ncohorts (P/R/F1 micro; single-VS -> watch RECALL, precision is count-capped):")
    single_recall = 0.0
    multi_f1 = _div(2 * micro_p * micro_r, micro_p + micro_r)
    for label, crows in cohorts:
        if not crows:
            print(f"  {label}  n=0  (none in this run - use --min-gt 1 to include single-VS)")
            continue
        cp, cr, cf = _cohort_prf(crows)
        print(f"  {label}  n={len(crows):3}  P={cp:.3f} R={cr:.3f} F1={cf:.3f}")
        if label.startswith("single-VS"):
            single_recall = cr
        if label.startswith("multi-VS"):
            multi_f1 = cf

    cohort_rows = {label: _cohort_prf(crows) for label, crows in cohorts if crows}
    cohort_n = {label: len(crows) for label, crows in cohorts}

    times = sorted(r["seconds"] for r in rows if r["seconds"] > 0)
    latency = {}
    if times:
        slow = min(rows, key=lambda r: -r["seconds"])
        fast = min(rows, key=lambda r: r["seconds"] if r["seconds"] > 0 else 1e9)
        retr = [r["retrieval_s"] for r in rows if r.get("retrieval_s")]
        sel = [r["selection_s"] for r in rows if r.get("selection_s")]
        latency = {"avg": _div(sum(times), len(times)), "min": fast["seconds"],
                   "max": slow["seconds"], "median": times[len(times) // 2],
                   "retrieval_avg": _div(sum(retr), len(retr)) if retr else 0.0,
                   "selection_avg": _div(sum(sel), len(sel)) if sel else 0.0}
        print(f"\nprediction latency (per ticket, excludes judge/explain):")
        print(f"  combined  avg={latency['avg']:.1f}s   min={latency['min']:.1f}s ({fast['ticket_id']})   "
              f"max={latency['max']:.1f}s ({slow['ticket_id']})   median={latency['median']:.1f}s")
        print(f"  split     retrieval avg={latency['retrieval_avg']:.2f}s   "
              f"LLM-selection avg={latency['selection_avg']:.1f}s  (= where the time goes)")

    if judge_pred:
        # Judge precision: predictions the LLM judge calls relevant (credits relevant non-GT picks).
        # Judge recall: TP / (TP + misses the judge says are genuinely supported) - drops GT label noise.
        adj_p = _div(judge_rel_pred, judge_pred)
        adj_r = _div(micro_tp, micro_tp + judge_supported_miss)
        print(f"judge  P={adj_p:.3f}  R={adj_r:.3f}  F1={_div(2*adj_p*adj_r, adj_p+adj_r):.3f}  "
              f"(LLM relevance; {judge_rel_pred}/{judge_pred} preds relevant, "
              f"{judge_supported_miss} of {micro_fn} misses judged supported)")

    total_fn = sum(bucket_totals.values())
    print(f"\nmiss buckets (where the {total_fn} missed GT value streams died):")
    print(f"  not_retrieved  {bucket_totals['not_retrieved']:4}  "
          f"({_div(bucket_totals['not_retrieved'], total_fn):.0%})  - never made the candidate set")
    print(f"  gated_pre_llm  {bucket_totals['gated_pre_llm']:4}  "
          f"({_div(bucket_totals['gated_pre_llm'], total_fn):.0%})  - merger dropped before the LLM")
    print(f"  llm_dropped    {bucket_totals['llm_dropped']:4}  "
          f"({_div(bucket_totals['llm_dropped'], total_fn):.0%})  - LLM saw it, didn't pick it")

    if reason_totals:
        explained = sum(reason_totals.values())
        print(f"\nwhy the LLM dropped them (of {explained} llm_dropped explained):")
        for code, cnt in sorted(reason_totals.items(), key=lambda kv: -kv[1]):
            print(f"  {code:24} {cnt:4}  ({_div(cnt, explained):.0%})")
    elif args.explain_drops and bucket_totals["llm_dropped"]:
        print(f"\n[!] --explain-drops was on and {bucket_totals['llm_dropped']} GT were llm_dropped, "
              "but the probe returned no reasons - check the explain-drops failure lines above.")
    elif bucket_totals["llm_dropped"]:
        print(f"\n[i] {bucket_totals['llm_dropped']} GT were llm_dropped but not explained - "
              "re-run with --explain-drops to classify why.")

    if swap_totals:  # Level C
        sw = sum(swap_totals.values())
        print(f"\n[Level C] why the picks beat each dropped GT (of {sw} explained):")
        for code, cnt in sorted(swap_totals.items(), key=lambda kv: -kv[1]):
            print(f"  {code:34} {cnt:4}  ({_div(cnt, sw):.0%})")
        real = swap_totals.get("dropped_is_valid_should_have_picked", 0)
        print(f"  -> {_div(real, sw):.0%} are real misses (model agrees the dropped GT was applicable)")

    if margins:  # Level B
        margins_sorted = sorted(margins)
        near = sum(1 for m in margins if m >= 0)  # dropped GT scored >= the pick cut
        close = sum(1 for m in margins if -0.1 <= m < 0)
        print(f"\n[Level B] dropped-GT score margin to the pick cut ({len(margins)} drops):")
        print(f"  near-miss (margin>=0, scored >= a pick): {near:4}  ({_div(near, len(margins)):.0%})  "
              "- the selection contradicted its own scoring - PROMPT-FIXABLE")
        print(f"  close     (-0.1..0):                     {close:4}  ({_div(close, len(margins)):.0%})")
        print(f"  median margin={margins_sorted[len(margins_sorted)//2]:+.2f}  "
              f"min={margins_sorted[0]:+.2f}  max={margins_sorted[-1]:+.2f}")

    out = Path(args.out)
    out.parent.mkdir(parents=True, exist_ok=True)
    with out.open("w", newline="", encoding="utf-8") as fh:
        w = csv.DictWriter(fh, fieldnames=list(rows[0].keys()))
        w.writeheader(); w.writerows(rows)
    print(f"\nper-ticket CSV -> {out}")

    return {
        "n": n, "micro_p": micro_p, "micro_r": micro_r,
        "micro_f1": _div(2 * micro_p * micro_r, micro_p + micro_r),
        "macro_f1": _div(2 * macro_p * macro_r, macro_p + macro_r),
        "single_recall": single_recall, "multi_f1": multi_f1,
        "avg_predicted": avg_pred, "avg_gt": avg_gt, "cardinality": cardinality,
        "count_following": _count_following_stats(rows),
        "cohorts": cohort_rows, "cohort_n": cohort_n, "latency": latency,
        "buckets": dict(bucket_totals), "judge_p": _div(judge_rel_pred, judge_pred) if judge_pred else None,
        "retrieval": retrieval_mean,
        "boost": {"backed_recall": br, "notbacked_recall": nr_, "lift": br - nr_},
        "drop_reasons": dict(reason_totals), "swap_reasons": dict(swap_totals),
        "score_margin_near_miss": _div(sum(1 for m in margins if m >= 0), len(margins)) if margins else None,
    }


def _mean_std(values: list[float]) -> tuple[float, float]:
    if not values:
        return 0.0, 0.0
    mean = sum(values) / len(values)
    var = sum((v - mean) ** 2 for v in values) / len(values)
    return mean, var ** 0.5


def build_eval_docx(args, runs: list[dict], out_path: Path) -> None:
    from datetime import date

    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    def table(doc, headers, rows):
        t = doc.add_table(rows=1, cols=len(headers)); t.style = "Light Grid Accent 1"
        for c, h in zip(t.rows[0].cells, headers):
            c.text = h
        for r in rows:
            cells = t.add_row().cells
            for c, v in zip(cells, r):
                c.text = str(v)
        doc.add_paragraph()

    doc = Document()
    title = doc.add_heading(f"VS Selection Eval — mode: {args.mode}", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph(
        f"{args.repeat} run(s) · input={'rawText' if args.raw_text else 'summary'} · "
        f"count={args.count} · min_gt={args.min_gt} · {date.today().isoformat()}")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in sub.runs:
        run.font.size = Pt(11); run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_heading("Headline (mean ± std over runs)", level=1)
    rows = []
    for key, lbl in [("micro_f1", "micro F1"), ("micro_p", "micro P"), ("micro_r", "micro R"),
                     ("macro_f1", "macro F1"), ("single_recall", "single-VS recall"),
                     ("multi_f1", "multi-VS F1")]:
        m, s = _mean_std([r[key] for r in runs])
        rows.append([lbl, f"{m:.3f} ± {s:.3f}", str([round(r[key], 3) for r in runs])])
    table(doc, ["Metric", "Mean ± Std", "Per run"], rows)

    nr = len(runs)
    avg = lambda key: sum(r.get(key, 0) or 0 for r in runs) / nr  # noqa: E731

    doc.add_heading(f"Cohorts (mean of {nr} runs)", level=1)
    doc.add_paragraph("Single-VS: watch recall (precision is count-capped). Multi-VS is the hard half.")
    crows = []
    for label in runs[0].get("cohorts", {}):
        ps = [r["cohorts"][label] for r in runs if label in r.get("cohorts", {})]
        if not ps:
            continue
        p = sum(x[0] for x in ps) / len(ps); r_ = sum(x[1] for x in ps) / len(ps); f = sum(x[2] for x in ps) / len(ps)
        crows.append([label.strip(), runs[0]["cohort_n"].get(label, "-"), f"{p:.3f}", f"{r_:.3f}", f"{f:.3f}"])
    table(doc, ["Cohort", "n", "P", "R", "F1"], crows)

    doc.add_heading(f"Retrieval recall (ceiling, mean of {nr} runs)", level=1)
    doc.add_paragraph("Did retrieval put the GT value streams in front of the LLM? The selection "
                      "recall cannot exceed the review-pool recall.")
    keys = list(runs[0].get("retrieval", {}))
    rrows = [[k.replace("vs_lane@", "VS lane R@").replace("historic_lane", "historic lane R")
              .replace("pool", "review pool R (ceiling)"),
              f"{sum(r['retrieval'].get(k, 0) for r in runs) / nr:.3f}"] for k in keys]
    table(doc, ["Lane / stage", "Recall (mean)"], rrows)

    if any(r.get("boost") for r in runs):
        doc.add_heading(f"Historic boost (mean of {nr} runs)", level=1)
        doc.add_paragraph("Recall on GT that the historic evidence surfaced vs GT it didn't. The "
                          "lift is how much the historic signal improves selection.")
        b0 = lambda k: sum((r.get("boost") or {}).get(k, 0) for r in runs) / nr  # noqa: E731
        table(doc, ["Metric", "Recall"], [
            ["GT backed by historic", f"{b0('backed_recall'):.3f}"],
            ["GT not in historic", f"{b0('notbacked_recall'):.3f}"],
            ["lift from historic", f"{b0('lift'):+.3f}"],
        ])

    doc.add_heading(f"Latency + pool (mean of {nr} runs)", level=1)
    bk = lambda k: sum((r.get("buckets") or {}).get(k, 0) for r in runs) / nr  # noqa: E731
    jp = [r["judge_p"] for r in runs if r.get("judge_p") is not None]
    table(doc, ["Metric", "Value (mean)"], [
        ["avg latency (s)", f"{sum(r.get('latency', {}).get('avg', 0) for r in runs)/nr:.1f}"],
        ["median latency (s)", f"{sum(r.get('latency', {}).get('median', 0) for r in runs)/nr:.1f}"],
        ["max latency (s)", f"{max(r.get('latency', {}).get('max', 0) for r in runs):.1f}"],
        ["avg predicted", f"{avg('avg_predicted'):.1f}"],
        ["avg gt", f"{avg('avg_gt'):.1f}"],
        ["judge precision", f"{sum(jp)/len(jp):.3f}" if jp else "n/a"],
        ["miss: gated_pre_llm", f"{bk('gated_pre_llm'):.0f}"],
        ["miss: llm_dropped", f"{bk('llm_dropped'):.0f}"],
        ["miss: not_retrieved", f"{bk('not_retrieved'):.0f}"],
    ])
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))


async def run_repeats(args) -> None:
    """Run the eval --repeat times and report each run + mean ± std (captures LLM variance)."""
    if args.rebuild_docx:
        # Build the docx from a saved <out>.runs.json without re-running the eval.
        runs = json.loads(Path(args.rebuild_docx).read_text(encoding="utf-8"))
        args.repeat = len(runs)
        docx_path = Path(args.out).with_name(f"eval_{args.mode}_{'raw' if args.raw_text else 'summary'}.docx")
        build_eval_docx(args, runs, docx_path)
        print(f"rebuilt docx from {args.rebuild_docx} -> {docx_path}")
        return
    runs = []
    for i in range(1, args.repeat + 1):
        print("\n" + "#" * 60 + f"\n# RUN {i}/{args.repeat}\n" + "#" * 60)
        runs.append(await main(args))
    if args.repeat > 1:
        print("\n" + "=" * 60 + f"\nSUMMARY over {args.repeat} runs (mean ± std):")
        for key, label in [("micro_p", "micro P"), ("micro_r", "micro R"),
                           ("micro_f1", "micro F1"), ("macro_f1", "macro F1"),
                           ("single_recall", "single-VS R"), ("multi_f1", "multi-VS F1")]:
            m, s = _mean_std([r[key] for r in runs])
            print(f"  {label:12} {m:.3f} ± {s:.3f}   runs={[round(r[key], 3) for r in runs]}")
    # Always dump the run metrics so the docx can be rebuilt without re-running the eval.
    stats_path = Path(args.out).with_suffix(".runs.json")
    stats_path.write_text(json.dumps([{k: v for k, v in r.items() if k != "cohorts" or True}
                                      for r in runs], indent=2, default=list), encoding="utf-8")
    print(f"run metrics -> {stats_path}")
    if args.docx:
        tag = "_" + args.selection_prompt.split("/")[-1] if args.selection_prompt else ""
        docx_path = Path(args.out).with_name(
            f"eval_{args.mode}{tag}_{'raw' if args.raw_text else 'summary'}.docx")
        try:
            build_eval_docx(args, runs, docx_path)
            print(f"\ndocx -> {docx_path}")
        except ImportError:
            print("\n[!] docx skipped: python-docx not installed. Run: uv sync --extra extract"
                  f"\n    (run metrics saved at {stats_path}; rebuild later with --rebuild-docx)")
        except PermissionError:
            print(f"\n[!] docx skipped: {docx_path} is locked (open in Word?). Close it, then:"
                  f"\n    uv run python -m scripts.eval_vs {args.dataset} --mode {args.mode} "
                  f"--rebuild-docx {stats_path} --out {args.out}")


if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("dataset", help="cosmos_idmt.json (ingested IDMT docs with GT)")
    parser.add_argument("--count", type=int, default=10, help="value streams to request (fixed mode)")
    parser.add_argument("--count-mode", choices=["fixed", "gt", "gt_buffer"], default="fixed",
                        help="fixed=--count; gt=|GT| per ticket (R-precision); gt_buffer=|GT|+buffer")
    parser.add_argument("--buffer", type=int, default=2, help="added to |GT| in gt_buffer mode")
    parser.add_argument("--min-gt", type=int, default=2, help="skip tickets with fewer than this many GT value streams")
    parser.add_argument("--k", type=int, nargs="+", default=[3, 5, 10], help="k values for P@k / R@k")
    parser.add_argument("--historic-k", type=int, default=0,
                        help="how many similar past tickets to retrieve/show as evidence "
                             "(overrides the config default of 6; e.g. 8 or 10)")
    parser.add_argument("--historic-repr", default="snippet",
                        choices=["snippet", "summary", "description", "raw"],
                        help="how to render each historic ticket in the evidence block (experiment)")
    parser.add_argument("--historic-budget", type=int, default=0,
                        help="truncate each historic 'raw' to ~N tokens (the K sweep; 0 = no cap)")
    parser.add_argument("--query-budget", type=int, default=0,
                        help="with --raw-text, truncate the query raw text to ~N tokens (e.g. 7000)")
    parser.add_argument("--concurrency", type=int, default=3, help="tickets evaluated in parallel")
    parser.add_argument("--semantic-only", action="store_true", help="ablation: drop the historic lane entirely")
    parser.add_argument("--raw-text", action="store_true", help="use rawText instead of summaryFields")
    parser.add_argument("--raw-retrieval", action="store_true",
                        help="retrieve with rawText (truncated to --retrieval-budget) instead of the "
                             "summary - ONLY valid against a raw-embedded index (see reembed_index.py)")
    parser.add_argument("--retrieval-budget", type=int, default=7000,
                        help="token budget for the raw retrieval query (with --raw-retrieval)")
    parser.add_argument("--window", type=int, default=0,
                        help="override the LLM review-pool size (how many candidates the LLM sees; "
                             "default config=18). Decoupled from output count, so count=gt stays honest.")
    parser.add_argument("--generic-penalty", type=float, default=0.0,
                        help="broad-stream rank penalty scale (penalty = scale * signal, unless earned "
                             "by history). Default 0 (off - avoids the fp_rate pass-1). Pass 0.6 for the "
                             "production operating point. Signal is leave-one-out.")
    parser.add_argument("--min-confidence", type=float, default=0.0,
                        help="abstention floor (0-1): keep only picks at/above this confidence and "
                             "stop padding to --count. Raises precision. Try 0.45. 0 = pad to count.")
    parser.add_argument("--penalty-signal", choices=["fp_rate", "gt_freq"], default="fp_rate",
                        help="fp_rate = false-positive rate (correct attractor signal, runs a pass 1); "
                             "gt_freq = corpus GT frequency (penalizes common-true streams - usually wrong).")
    parser.add_argument("--explain-drops", action="store_true",
                        help="post-hoc LLM probe: classify why each llm_dropped GT was left out (extra calls)")
    parser.add_argument("--score-margins", action="store_true",
                        help="Level B: LLM scores every candidate 0-1; report how close dropped GT was "
                             "to the cut (near-miss vs genuine reject). Extra calls.")
    parser.add_argument("--explain-swaps", action="store_true",
                        help="Level C: comparative probe - why the picks beat each dropped GT, in a "
                             "richer taxonomy (more actionable than --explain-drops). Extra calls.")
    parser.add_argument("--judge", action="store_true",
                        help="LLM-as-judge relevance of predictions + misses (GT-independent view; extra calls)")
    parser.add_argument("--mode", choices=["merge", "all50", "topk", "historic_only", "evidence"],
                        default="merge",
                        help="candidate-pool strategy: merge (VS+historic), all50 (all VS, no historic), "
                             "topk (top-K VS only), historic_only (historic VS only), evidence (all VS + "
                             "historic as an evidence block, no merge)")
    parser.add_argument("--repeat", type=int, default=1,
                        help="run the eval N times and report mean +/- std (captures LLM variance)")
    parser.add_argument("--docx", action="store_true", help="write a .docx report of the run(s)")
    parser.add_argument("--rebuild-docx", default="",
                        help="build the docx from a saved <out>.runs.json WITHOUT re-running the eval")
    parser.add_argument("--selection-prompt", default="",
                        help="override the mode's selection prompt (A/B prompt variants), "
                             "e.g. value_stream/selection_plain_lean")
    parser.add_argument("--sample", type=int, default=0,
                        help="evaluate a seeded random subset of N tickets (fast prompt iteration); "
                             "0 = all. Same N+seed -> same tickets across runs, so A/Bs are comparable.")
    parser.add_argument("--seed", type=int, default=13, help="sampling seed (fixed so runs match)")
    parser.add_argument("--no-candidate-scores", action="store_true",
                        help="strip the lane + semantic score from candidate blocks (weak signal when "
                             "VS-lane recall is low - test if the model does better without it)")
    parser.add_argument("--attachments-cache", default="",
                        help="EDA attachments_raw.json - adds no-attachment cohort rows to the breakdown")
    parser.add_argument("--out", default="out/eval/vs_eval.csv")
    args = parser.parse_args()
    asyncio.run(run_repeats(args))
