"""Authored, charted retrieval-findings report (docx) - plain language, real charts.

Numbers are from the 373-ticket historic-lane retrieval run (data in DATA below; cohort +
gt-distribution from retrieval_cohort.py / gt_distribution.py). The prose and chart choices are
authored for a non-DS reader. One command, no flags:

    uv sync --extra eda --extra extract
    uv run python scripts/build_retrieval_report.py   ->  out/eval/retrieval_report.docx
"""

from __future__ import annotations

from pathlib import Path

OUT = Path("out/eval/retrieval_report.docx")

# ---- real numbers from the 373-ticket run ----------------------------------------------------
DATA = {
    "n": 373,
    "k": [6, 8, 10],
    "recall": [0.902, 0.925, 0.944],          # avg % of GT covered
    "precision": [0.672, 0.647, 0.631],       # shares-a-tag precision
    "precision_strict": [0.333, 0.314, 0.301],  # minus broad tags
    "content_precision": [0.365, 0.337, 0.323],  # AI judge
    "hit": [0.960, 0.968, 0.976],
    "fully_covered": [0.780, 0.812, 0.850],
    "mrr": [0.846, 0.848, 0.848],
    # label vs content over 3,729 judged retrieved tickets
    "crosstab": {"real": 914, "lucky": 1438, "mislabeled": 292, "unrelated": 1085, "total": 3729},
    # single vs multi coverage at K=6/8/10
    "cov_single": [0.922, 0.938, 0.953],
    "cov_multi": [0.880, 0.910, 0.934],
    "full_single": [0.922, 0.938, 0.953],
    "full_multi": [0.628, 0.678, 0.739],
    "single_n": 193, "multi_n": 180,
    # GT size buckets
    "gt_buckets": {"Just 1": 193, "2-4": 98, "5-9": 54, "10+": 28},
    # first-relevant rank histogram
    "first_rank": {"1": 288, "2": 36, "3": 20, "4-10": 20, "none": 9},
    # tags per pulled ticket (evidence density)
    "density": {"1": 1885, "2-5": 1047, "6-10": 496, "11-18": 209, "19": 92},
    "score_rel": 0.510, "score_irr": 0.480,
}


def _save(fig, plt, charts, name):
    p = charts / f"{name}.png"
    fig.savefig(p, dpi=150, bbox_inches="tight"); plt.close(fig)
    return str(p)


def _bars(ax, labels, values, colors=None, pct=True, rotate=0):
    bars = ax.bar(labels, values, color=colors)
    for b, v in zip(bars, values):
        ax.text(b.get_x() + b.get_width() / 2, v, f"{v:.0%}" if pct else f"{v:,}",
                ha="center", va="bottom", fontsize=9, fontweight="bold")
    if rotate:
        ax.set_xticklabels(labels, rotation=rotate, ha="right")
    ax.grid(axis="y", alpha=0.3)


def build() -> None:
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor

    d = DATA
    charts = OUT.parent / "retrieval_report_charts"
    charts.mkdir(parents=True, exist_ok=True)
    GREEN, AMBER, BLUE, GREY = "#2a9d4a", "#e8a23d", "#3d7fe8", "#9aa0a6"

    # 1. coverage scorecard
    fig, ax = plt.subplots(figsize=(7.5, 4))
    _bars(ax, ["Right answer\nfound (@6)", "Found ≥1 useful\nexample (@6)", "Found EVERY\nstream (@6)"],
          [d["recall"][0], d["hit"][0], d["fully_covered"][0]], colors=[GREEN, GREEN, BLUE])
    ax.set_ylim(0, 1.1); ax.figure.suptitle("The right examples almost always show up", fontsize=14, fontweight="bold")
    ax.set_title("Coverage is the strong part of the system.", fontsize=10, style="italic", color="#555")
    c1 = _save(fig, plt, charts, "coverage")

    # 2. precision reality check
    fig, ax = plt.subplots(figsize=(7.5, 4))
    _bars(ax, ["Looks relevant\n(shares a tag)", "Really relevant\n(minus generic tags)",
               "Really relevant\n(AI double-check)"],
          [d["precision"][0], d["precision_strict"][0], d["content_precision"][0]], colors=[AMBER, GREEN, GREEN])
    ax.set_ylim(0, 0.8); ax.figure.suptitle("But about half the 'matches' aren't really relevant", fontsize=14, fontweight="bold")
    ax.set_title("Two independent checks agree real relevance is ~33-37%, not 67%.", fontsize=10, style="italic", color="#555")
    c2 = _save(fig, plt, charts, "precision_check")

    # 3. cross-tab breakdown
    ct = d["crosstab"]; tot = ct["total"]
    fig, ax = plt.subplots(figsize=(7.5, 4))
    _bars(ax, ["Real match", "Lucky match\n(coincidence)", "Same work,\nwrong tag", "Unrelated"],
          [ct["real"] / tot, ct["lucky"] / tot, ct["mislabeled"] / tot, ct["unrelated"] / tot],
          colors=[GREEN, AMBER, BLUE, GREY])
    ax.set_ylim(0, 0.55); ax.figure.suptitle("Every example we pulled, sorted", fontsize=14, fontweight="bold")
    ax.set_title(f"Of {tot:,} pulled tickets: 6 in 10 'matches' are lucky coincidences.", fontsize=10, style="italic", color="#555")
    c3 = _save(fig, plt, charts, "crosstab")

    # 4. recall vs precision across K
    fig, ax = plt.subplots(figsize=(7.5, 4))
    ax.plot(d["k"], d["recall"], "-o", color=GREEN, label="Coverage (right answers found)")
    ax.plot(d["k"], d["precision"], "-s", color=AMBER, label="Looks relevant")
    ax.plot(d["k"], d["content_precision"], "--^", color=BLUE, label="Really relevant (AI)")
    for kk, r, p in zip(d["k"], d["recall"], d["precision"]):
        ax.annotate(f"{r:.0%}", (kk, r), textcoords="offset points", xytext=(0, 7), ha="center", fontsize=8)
        ax.annotate(f"{p:.0%}", (kk, p), textcoords="offset points", xytext=(0, -14), ha="center", fontsize=8)
    ax.set_xticks(d["k"]); ax.set_xlabel("Number of past tickets shown"); ax.set_ylim(0, 1.05)
    ax.figure.suptitle("More examples = more coverage, but less relevance", fontsize=14, fontweight="bold")
    ax.set_title("Adding tickets trades coverage for relevance ~1-for-1. 6 is the sweet spot.", fontsize=10, style="italic", color="#555")
    ax.legend(fontsize=8); ax.grid(alpha=0.3)
    c4 = _save(fig, plt, charts, "tradeoff")

    # 5. easy vs hard coverage
    import numpy as np
    fig, ax = plt.subplots(figsize=(7.5, 4))
    x = np.arange(2); w = 0.35
    avg = [d["cov_single"][0], d["cov_multi"][0]]
    full = [d["full_single"][0], d["full_multi"][0]]
    b1 = ax.bar(x - w / 2, avg, w, label="Found MOST streams (avg)", color=BLUE)
    b2 = ax.bar(x + w / 2, full, w, label="Found EVERY stream", color=GREEN)
    for bars in (b1, b2):
        for b in bars:
            ax.text(b.get_x() + b.get_width() / 2, b.get_height(), f"{b.get_height():.0%}",
                    ha="center", va="bottom", fontsize=9, fontweight="bold")
    ax.set_xticks(x); ax.set_xticklabels([f"Easy\n(1 stream, {d['single_n']})", f"Hard\n(2+, {d['multi_n']})"])
    ax.set_ylim(0, 1.1); ax.figure.suptitle("Hard tickets: we find most, but rarely all", fontsize=14, fontweight="bold")
    ax.set_title("On hard tickets, avg coverage is 88% but full coverage drops to 63%.", fontsize=10, style="italic", color="#555")
    ax.legend(fontsize=8); ax.grid(axis="y", alpha=0.3)
    c5 = _save(fig, plt, charts, "easy_hard")

    # 6. GT distribution
    fig, ax = plt.subplots(figsize=(7.5, 4))
    _bars(ax, list(d["gt_buckets"]), list(d["gt_buckets"].values()),
          colors=[GREEN, BLUE, AMBER, "#c0392b"], pct=False)
    ax.figure.suptitle("How many Value Streams does a ticket have?", fontsize=14, fontweight="bold")
    ax.set_title("Half the tickets have just one (easy); a tail have up to 19 (hard).", fontsize=10, style="italic", color="#555")
    c6 = _save(fig, plt, charts, "gt_dist")

    # 7. first relevant rank
    fig, ax = plt.subplots(figsize=(7.5, 4))
    fr = d["first_rank"]
    _bars(ax, list(fr), list(fr.values()), colors=[GREEN, BLUE, BLUE, BLUE, GREY], pct=False)
    ax.figure.suptitle("Where the first useful example lands", fontsize=14, fontweight="bold")
    ax.set_title("For 288 of 373 tickets the very first result is already useful.", fontsize=10, style="italic", color="#555")
    ax.set_xlabel("position in the list")
    c7 = _save(fig, plt, charts, "first_rank")

    # 8. evidence density
    fig, ax = plt.subplots(figsize=(7.5, 4))
    de = d["density"]
    _bars(ax, list(de), list(de.values()), colors=[GREEN, BLUE, AMBER, AMBER, "#c0392b"], pct=False)
    ax.figure.suptitle("How many tags each pulled ticket carries", fontsize=14, fontweight="bold")
    ax.set_title("Most are specific (1 tag); the '19-tag' overloaded tickets cause the lucky matches.", fontsize=10, style="italic", color="#555")
    ax.set_xlabel("tags on the ticket")
    c8 = _save(fig, plt, charts, "density")

    # ---------------- document
    doc = Document()
    doc.add_heading("How good is our “similar past tickets” search?", level=0)
    doc.add_paragraph(
        "A plain-language report on the search that finds similar past tickets and shows them to the "
        "model as examples. Tested on 373 real tickets. Each chart says what it shows and how to read "
        "it; a glossary is at the end.")

    doc.add_heading("Bottom line", level=1)
    for b in [
        "The search is GREAT at finding the right examples — the correct Value Stream shows up for "
        "90% of tickets, usually as the very first result.",
        "But about HALF of what it calls a “match” only matches by coincidence — two "
        "independent checks (ignoring generic tags, and an AI reviewer) both say real relevance is ~33–37%, "
        "not the 67% the simple count suggests.",
        "The cause: 6 generic catch-all tags + a tail of “overloaded” tickets tagged with up to 19 "
        "Value Streams that match almost anything.",
        "Showing 6 past tickets is the sweet spot — more dilutes the examples without really helping.",
        "Fix worth doing: down-weight the generic tags and overloaded tickets when ranking — it sharpens "
        "relevance AND helps cover the hard multi-stream tickets.",
    ]:
        doc.add_paragraph(b, style="List Bullet")

    sections = [
        ("1. It finds the right examples", c1,
         "For 90% of tickets the correct Value Stream is among the 6 pulled examples, 96% have at least "
         "one genuinely useful example, and for 288 of 373 tickets the FIRST result is already a hit. "
         "Finding good examples is not the problem."),
        ("2. Not every ticket is equally hard", c6,
         "Half the tickets belong to a single Value Stream (easy — one answer to find); the other half "
         "belong to several, a few to as many as 19 (hard). So the coverage numbers are an average across "
         "easy and hard tickets."),
        ("3. Hard tickets: we find most, but rarely all", c5,
         "On hard (multi-stream) tickets the search still surfaces 88% of the streams on average — almost "
         "as good as the 92% for easy tickets. But getting EVERY stream right happens only 63% of the time "
         "(vs 92% for easy). In plain terms: on a hard ticket we catch the obvious streams and usually miss "
         "one or two from the long tail."),
        ("4. But “relevant” is overcounted", c2,
         "If we just count “shares a tag”, 67% of pulled tickets look relevant. That is misleading: "
         "removing the 6 generic tags drops it to 33%, and an AI reviewer reading the actual text puts it at "
         "37%. A simple rule and an AI independently land in the same place — real relevance is about half "
         "the headline."),
        ("5. Every example, sorted", c3,
         "Splitting all 3,729 pulled tickets four ways: only 24% are real matches (right tag AND same work), "
         "39% are lucky matches (right tag, different work), 8% are the same work tagged differently, and 29% "
         "are unrelated. So 6 in 10 “matches” are coincidences."),
        ("6. Why the lucky matches happen", c8,
         "Most pulled tickets are specific (a single tag). But a tail of “overloaded” tickets carry up "
         "to 19 tags each — those match almost any query by accident. Combined with 6 generic catch-all tags, "
         "they create the coincidental matches."),
        ("7. More examples is not better", c4,
         "Showing 6, 8, or 10 past tickets: more finds slightly more right answers but a steady share of the "
         "extras are junk — a near 1-for-1 trade. 6 is the sweet spot."),
        ("8. The first useful example is usually right at the top", c7,
         "For 288 of 373 tickets the #1 result is already useful; only 9 tickets found nothing relevant at "
         "all. The ranking puts good examples first — though the underlying scores barely separate good "
         "from bad (0.51 vs 0.48), so we rely on the ordering, not the score size."),
    ]
    for title, chart, text in sections:
        doc.add_heading(title, level=1)
        doc.add_paragraph(text)
        doc.add_picture(chart, width=Inches(6.0))

    doc.add_heading("All the numbers", level=1)
    t = doc.add_table(rows=1, cols=4); t.style = "Light Grid Accent 1"
    for c, h in zip(t.rows[0].cells, ["Measure (per ticket)", "Show 6", "Show 8", "Show 10"]):
        c.text = h
    rows = [("Right answer found (coverage)", "recall"), ("Found at least one useful", "hit"),
            ("Found EVERY stream", "fully_covered"), ("Looks relevant", "precision"),
            ("Really relevant (minus generic)", "precision_strict"), ("Really relevant (AI)", "content_precision"),
            ("First useful at the top (MRR)", "mrr")]
    for label, key in rows:
        cells = t.add_row().cells
        cells[0].text = label
        for i, v in enumerate(d[key]):
            cells[i + 1].text = f"{v:.0%}" if key != "mrr" else f"{v:.2f}"

    doc.add_heading("Plain-language glossary", level=1)
    gl = doc.add_table(rows=1, cols=2); gl.style = "Light Grid Accent 1"
    gl.rows[0].cells[0].text, gl.rows[0].cells[1].text = "Term", "What it means"
    for term, mean in [
        ("Value Stream / tag", "The business category a ticket belongs to — what we predict."),
        ("Coverage", "Did the correct Value Stream show up among the pulled examples."),
        ("Looks relevant", "An example that shares a tag with the current ticket."),
        ("Really relevant", "An example that's actually about the same kind of work, not just a shared tag."),
        ("Lucky match", "Shares a tag by coincidence (usually a generic tag) but is different work."),
        ("Generic / catch-all tag", "A Value Stream on a large share of tickets, so sharing it means little."),
        ("Overloaded ticket", "A past ticket tagged with many Value Streams — matches almost anything."),
        ("Easy / hard ticket", "Easy = 1 correct Value Stream; hard = 2 or more."),
    ]:
        cells = gl.add_row().cells
        cells[0].text, cells[1].text = term, mean

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print(f"report -> {OUT}")


if __name__ == "__main__":
    build()
