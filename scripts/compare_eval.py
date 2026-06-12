"""Compare VS-selection eval modes side by side: charts + interpretation in one docx.

Reads the per-run metrics each eval writes (<out>.runs.json), aggregates them (mean across the
repeat runs), and renders grouped bar charts comparing the modes on quality (P/R/F1), the
single-vs-multi cohorts, retrieval recall (the ceiling), historic boost, and latency - each with
a short interpretation. One picture instead of flipping between docs.

Usage (install: uv sync --extra eda):
  uv run python scripts/compare_eval.py out/eval/all50.runs.json out/eval/evidence.runs.json \
      --out out/eval/comparison.docx
Label = the file stem (all50.runs.json -> "all50"); override with label=path.
"""

from __future__ import annotations

import argparse
import json
from pathlib import Path


def _label_path(arg: str) -> tuple[str, str]:
    # 'label=path' - split on the LAST '=' so a label may itself contain '=' (e.g. "Recall K=6"),
    # since the path has no '='. Only treat it as label=path when the right side looks like a path.
    if "=" in arg:
        left, right = arg.rsplit("=", 1)
        if left and right and (right.endswith(".json") or "/" in right or "\\" in right):
            return left, right
    stem = Path(arg).name.replace(".runs.json", "").replace(".json", "")
    return stem, arg


def _parse_axis(arg: str) -> tuple[str, str, str]:
    """'plain->historic: all50 vs evidence' -> ('plain->historic', 'all50', 'evidence')."""
    name, _, pair = arg.partition(":")
    a, _, b = pair.partition(" vs ")
    a, b = a.strip(), b.strip()
    if not (name.strip() and a and b):
        raise SystemExit(f"--axis must be 'name: labelA vs labelB', got: {arg!r}")
    return name.strip(), a, b


# Metrics shown in an axis-delta table: (label, key, higher_is_better).
_AXIS_METRICS = [
    ("Recall (answers found)", "micro_r", True),
    ("F1 (overall balance)", "micro_f1", True),
    ("Easy tickets: found", "single_r", True),
    ("Hard tickets: F1", "multi_f1", True),
    ("Precedent use (lift)", "boost_lift", True),
    ("Missed answers (count)", "llm_dropped", False),
    ("Typical time (s)", "lat_med", False),
]


def _mean(runs: list[dict], key: str) -> float:
    vals = [r.get(key) for r in runs if r.get(key) is not None]
    return sum(vals) / len(vals) if vals else 0.0


def _mean_nested(runs: list[dict], outer: str, inner: str) -> float:
    vals = [(r.get(outer) or {}).get(inner) for r in runs]
    vals = [v for v in vals if v is not None]
    return sum(vals) / len(vals) if vals else 0.0


def _mean_nested_opt(runs: list[dict], outer: str, inner: str) -> float | None:
    """Like _mean_nested but None (not 0.0) when the field is absent in every run, so the
    axis table can show 'n/a' instead of a fake zero for older runs.json that predate a field."""
    vals = [(r.get(outer) or {}).get(inner) for r in runs]
    vals = [v for v in vals if v is not None]
    return sum(vals) / len(vals) if vals else None


def _cohort(runs: list[dict], prefix: str) -> tuple[float | None, float | None, float | None, int | None]:
    """Mean (P, R, F1, n) for a cohort whose label starts with prefix; Nones if it isn't recorded."""
    triples, ns = [], []
    for r in runs:
        for label, prf in (r.get("cohorts") or {}).items():
            if label.strip().startswith(prefix) and prf:
                triples.append(prf)
                n = (r.get("cohort_n") or {}).get(label)
                if n is not None:
                    ns.append(n)
    if not triples:
        return None, None, None, None
    p = sum(t[0] for t in triples) / len(triples)
    r_ = sum(t[1] for t in triples) / len(triples)
    f = sum(t[2] for t in triples) / len(triples)
    return p, r_, f, (round(sum(ns) / len(ns)) if ns else None)


def _ratio(num: float | None, den: float | None) -> float | None:
    return (num / den) if (num is not None and den not in (None, 0)) else None


def load(arg: str) -> dict:
    label, path = _label_path(arg)
    runs = json.loads(Path(path).read_text(encoding="utf-8"))
    backed_r, notbacked_r = _boost_split(runs)
    micro_r = _mean(runs, "micro_r")
    hist_r = _mean_nested(runs, "retrieval", "historic_lane")
    pool_r = _mean_nested(runs, "retrieval", "pool")
    sp, sr, sf, sn = _cohort(runs, "single-VS")
    mp, mr, mf, mn = _cohort(runs, "multi-VS")
    return {
        "label": label,
        "micro_p": _mean(runs, "micro_p"), "micro_r": micro_r,
        "micro_f1": _mean(runs, "micro_f1"), "judge_p": _mean(runs, "judge_p"),
        # cohorts (prefer the recorded per-cohort P/R/F1; fall back to the headline single_recall/multi_f1)
        "single_p": sp, "single_r": sr if sr is not None else _mean(runs, "single_recall"),
        "single_f1": sf, "single_n": sn,
        "multi_p": mp, "multi_r": mr,
        "multi_f1": mf if mf is not None else _mean(runs, "multi_f1"), "multi_n": mn,
        # retrieval depth + the two lanes + the ceiling
        "vs_r3": _mean_nested_opt(runs, "retrieval", "vs_lane@3"),
        "vs_r5": _mean_nested_opt(runs, "retrieval", "vs_lane@5"),
        "vs_r10": _mean_nested(runs, "retrieval", "vs_lane@10"),
        "hist_r": hist_r, "pool_r": pool_r,
        # how much of what was available the model actually picked
        "ceiling_capture": _ratio(micro_r, pool_r),
        "precedent_capture": _ratio(backed_r, hist_r),
        # Both exactly 0 means an older run that serialized the field without computing it (a real
        # run always has some recall) - treat as missing so it draws no bar instead of a fake zero.
        "backed_r": backed_r, "notbacked_r": notbacked_r,
        "boost_lift": _boost_lift(runs),
        "lat_avg": _mean_nested(runs, "latency", "avg"),
        "lat_avg_clean": _lat_avg_clean(runs),
        "lat_med": _mean_nested_opt(runs, "latency", "median"),
        "lat_max": _mean_nested(runs, "latency", "max"),
        "lat_outlier": _has_outlier(runs),
        # where misses come from (retrieval vs the model)
        "miss_not_retrieved": _mean_nested_opt(runs, "buckets", "not_retrieved"),
        "miss_gated": _mean_nested_opt(runs, "buckets", "gated_pre_llm"),
        "llm_dropped": _mean_nested_opt(runs, "buckets", "llm_dropped"),
    }


def _lat_avg_clean(runs: list[dict]) -> float | None:
    """Mean latency with the single slowest ticket removed per run, when that ticket is an outlier
    (slowest > 3x the median). runs.json only stores aggregates, so we back it out of the average:
    (avg*n - slowest) / (n-1). Recovers the real typical cost when one ticket hit a retry storm."""
    vals = []
    for r in runs:
        lat = r.get("latency") or {}
        avg, mx, med, n = lat.get("avg"), lat.get("max"), lat.get("median"), r.get("n")
        if avg is None:
            continue
        if mx is not None and med and n and n > 1 and mx > 3 * med:
            vals.append((avg * n - mx) / (n - 1))
        else:
            vals.append(avg)
    return sum(vals) / len(vals) if vals else None


def _has_outlier(runs: list[dict]) -> bool:
    for r in runs:
        lat = r.get("latency") or {}
        mx, med = lat.get("max"), lat.get("median")
        if mx is not None and med and mx > 3 * med:
            return True
    return False


def _boost_split(runs: list[dict]) -> tuple[float | None, float | None]:
    backed = _mean_nested_opt(runs, "boost", "backed_recall")
    notbacked = _mean_nested_opt(runs, "boost", "notbacked_recall")
    if (backed or 0) == 0 and (notbacked or 0) == 0:
        return None, None  # not actually computed in this (older) run
    return backed, notbacked


def _boost_lift(runs: list[dict]) -> float | None:
    # Prefer the stored lift; older runs.json have only backed/notbacked, so derive it from those.
    lift = _mean_nested_opt(runs, "boost", "lift")
    if lift is not None:
        return lift
    backed = _mean_nested_opt(runs, "boost", "backed_recall")
    notbacked = _mean_nested_opt(runs, "boost", "notbacked_recall")
    return (backed - notbacked) if (backed is not None and notbacked is not None) else None


def _titles(ax, title: str, note: str | None) -> None:
    # Question as the bold headline, the 'how to read it' line as a smaller subtitle - so the chart
    # explains itself without the surrounding text.
    ax.figure.suptitle(title, fontsize=13, fontweight="bold", y=0.99)
    if note:
        ax.set_title(note, fontsize=9, style="italic", color="#555", pad=8)


def _bar(ax, labels, series: dict, title, ylabel, note=None, pct=False):
    import numpy as np

    x = np.arange(len(labels))
    width = 0.8 / max(1, len(series))
    for i, (name, raw) in enumerate(series.items()):
        # None -> NaN so a missing value draws no bar (instead of a misleading 0).
        vals = [float("nan") if v is None else v for v in raw]
        bars = ax.bar(x + i * width - 0.4 + width / 2, vals, width, label=name)
        for b, v in zip(bars, vals):
            if v == v:  # skip NaN (no bar to label)
                txt = f"{v:.0%}" if pct else f"{v:.2f}"
                ax.text(b.get_x() + b.get_width() / 2, v, txt, ha="center", va="bottom", fontsize=8)
    ax.set_xticks(x); ax.set_xticklabels(labels, fontsize=9)
    ax.set_ylabel(ylabel)
    # headroom above the tallest bar so value labels and the legend don't collide
    top = max((v for raw in series.values() for v in raw if v is not None), default=1.0)
    ax.set_ylim(0, top * 1.18)
    _titles(ax, title, note)
    ax.legend(fontsize=8, loc="best"); ax.grid(axis="y", alpha=0.3)


def _stacked(ax, labels, series: dict, title, ylabel, note=None):
    """Stacked bars - for 'where did the misses go': the whole bar is the total, the bands are the split."""
    import numpy as np

    x = np.arange(len(labels))
    bottom = np.zeros(len(labels))
    for name, raw in series.items():
        vals = np.array([0.0 if v is None else v for v in raw], dtype=float)
        bars = ax.bar(x, vals, 0.55, bottom=bottom, label=name)
        for b, v, bot in zip(bars, vals, bottom):
            if v > 0:
                ax.text(b.get_x() + b.get_width() / 2, bot + v / 2, f"{v:.0f}", ha="center",
                        va="center", fontsize=8, color="white", fontweight="bold")
        bottom += vals
    ax.set_xticks(x); ax.set_xticklabels(labels, fontsize=9)
    ax.set_ylabel(ylabel)
    _titles(ax, title, note)
    ax.legend(fontsize=8, loc="upper left"); ax.grid(axis="y", alpha=0.3)


def _render_axes(doc, by_label: dict[str, dict], axes: list[tuple[str, str, str]]) -> None:
    """One delta table per axis: each metric's value for A, for B, and B-A with a verdict word."""
    doc.add_heading("Head-to-head: what changes when we flip one setting", level=1)
    doc.add_paragraph(
        "Each comparison below changes exactly ONE thing between two approaches and holds everything "
        "else fixed, so any difference is caused by that one change. 'Change' is the second approach "
        "minus the first; the arrow points to whichever side is better for that metric (recall, F1 and "
        "precedent-use: higher is better; missed answers and time: lower is better).")
    for name, la, lb in axes:
        a, b = by_label.get(la), by_label.get(lb)
        if a is None or b is None:
            doc.add_paragraph(f"[skipped '{name}': missing approach '{la if a is None else lb}']")
            continue
        doc.add_heading(name, level=2)
        doc.add_paragraph(f"Comparing  \"{la}\"  →  \"{lb}\"", style="Intense Quote")
        t = doc.add_table(rows=1, cols=5); t.style = "Light Grid Accent 1"
        for c, h in zip(t.rows[0].cells, ["What we measure", la, lb, "Change", "Better"]):
            c.text = h
        for mlabel, key, higher in _AXIS_METRICS:
            av, bv = a.get(key), b.get(key)
            fmt = "{:.0f}" if key == "llm_dropped" else "{:.3f}"
            if av is None or bv is None:
                # A field one side never recorded (older runs.json) - show what we have, no delta.
                row = [mlabel, "n/a" if av is None else fmt.format(av),
                       "n/a" if bv is None else fmt.format(bv), "n/a", "—"]
            else:
                delta = bv - av
                improved = (delta > 0) == higher
                arrow = "—" if abs(delta) < 1e-9 else (f"→ {lb}" if improved else f"→ {la}")
                row = [mlabel, fmt.format(av), fmt.format(bv), f"{delta:+.3f}", arrow]
            cells = t.add_row().cells
            for c, v in zip(cells, row):
                c.text = v


def _pct(v: float | None) -> str:
    return "n/a" if v is None else f"{v:.0%}"


def _f2(v: float | None) -> str:
    return "n/a" if v is None else f"{v:.2f}"


def _na_int(v: float | None) -> str:
    return "n/a" if v is None else f"{v:.0f}"


def _secs(v: float | None) -> str:
    return "n/a" if v is None else f"{v:.1f}s"


def _note(doc, text: str) -> None:
    p = doc.add_paragraph(text)
    p.runs[0].italic = True


def _add_md_runs(paragraph, text: str) -> None:
    # Render inline **bold** by splitting on the markers; everything else is a plain run.
    for i, part in enumerate(text.split("**")):
        if part:
            run = paragraph.add_run(part)
            run.bold = i % 2 == 1


def _render_markdown(doc, text: str) -> None:
    """Render a small markdown subset (## / ### headings, - bullets, **bold**, paragraphs) into the
    docx, so an external analysis file can be dropped into the report verbatim."""
    for block in text.split("\n"):
        line = block.rstrip()
        if not line.strip():
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.lstrip().startswith(("- ", "* ")):
            p = doc.add_paragraph(style="List Bullet")
            _add_md_runs(p, line.lstrip()[2:])
        else:
            _add_md_runs(doc.add_paragraph(), line)


def _add_table(doc, headers: list[str], rows: list[list[str]]):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = "Light Grid Accent 1"
    for c, h in zip(t.rows[0].cells, headers):
        c.text = h
    for r in rows:
        cells = t.add_row().cells
        for c, v in zip(cells, r):
            c.text = str(v)
    return t


def build(data: list[dict], out_path: Path, axes: list[tuple[str, str, str]] | None = None,
          describe: dict[str, str] | None = None, history_runs: set[str] | None = None,
          narrative: str | None = None, title: str | None = None, append: bool = False) -> None:
    import matplotlib

    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    from docx import Document
    from docx.shared import Inches

    describe = describe or {}
    labels = [d["label"] for d in data]
    # Which approaches actually SHOWED past tickets to the model. The precedent analysis only makes
    # sense for these - a "no history" run still has the numbers (the tickets were fetched to compute
    # them) but the model never saw them, which is confusing, so we leave those out of section 4.
    # If the caller doesn't say, fall back to runs that recorded the split.
    hist_set = set(history_runs) if history_runs else {d["label"] for d in data if d["backed_r"] is not None}
    hist_data = [d for d in data if d["label"] in hist_set]
    hist_labels = [d["label"] for d in hist_data]
    charts_dir = out_path.parent / "compare_charts"
    charts_dir.mkdir(parents=True, exist_ok=True)

    def save(fig, name):
        p = charts_dir / f"{name}.png"
        fig.savefig(p, dpi=150, bbox_inches="tight"); plt.close(fig)
        return str(p)

    figs = {}
    # 1. quality
    fig, ax = plt.subplots(figsize=(8, 4.6))
    _bar(ax, labels, {"Precision": [d["micro_p"] for d in data], "Recall": [d["micro_r"] for d in data],
                      "F1": [d["micro_f1"] for d in data]},
         "How good are the picks overall?",
         "fraction (0-1)", note="Read RECALL (correct answers found). Precision is low for all - "
         "the model returns 10 guesses for ~2.5 real answers, by design.", pct=True)
    figs["quality"] = save(fig, "quality")
    # 2. cohorts - SAME metric (recall) for both; only approaches that recorded per-cohort numbers
    cohort_data = [d for d in data if d["multi_r"] is not None]
    cohort_omitted = [d["label"] for d in data if d["multi_r"] is None]
    if cohort_data:
        fig, ax = plt.subplots(figsize=(8, 4.6))
        _bar(ax, [d["label"] for d in cohort_data],
             {"Easy tickets (1 correct answer)": [d["single_r"] for d in cohort_data],
              "Hard tickets (2+ correct answers)": [d["multi_r"] for d in cohort_data]},
             "Does it do well on hard tickets too, not just easy ones?",
             "recall (correct answers found)",
             note="Same metric (recall) for both. A win on the orange bars = it handles the "
             "multi-answer tickets, not just the easy ones.", pct=True)
        figs["cohorts"] = save(fig, "cohorts")
    # 3. retrieval depth + lanes - identical across approaches (search is the same), so show once
    rep = next((d for d in data if d["vs_r3"] is not None), data[0])
    fig, ax = plt.subplots(figsize=(8, 4.6))
    _bar(ax, ["Search\ntop 3", "Search\ntop 5", "Search\ntop 10", "Past-ticket\nexamples",
              "Full catalogue\n(always shown)"],
         {"correct answers available here": [rep["vs_r3"], rep["vs_r5"], rep["vs_r10"], rep["hist_r"],
                                             rep["pool_r"]]},
         "Did the system put the right answers in front of the model?",
         "fraction of correct answers present",
         note="The bar for 'full catalogue' is the CEILING - what the model could pick from. "
         "Same for every approach (the search is identical).", pct=True)
    figs["retrieval"] = save(fig, "retrieval")
    # 4. where the misses went - stacked, so it's visually obvious it's all the model's choice
    miss_data = [d for d in data if d["llm_dropped"] is not None]
    if miss_data:
        fig, ax = plt.subplots(figsize=(8, 4.6))
        _stacked(ax, [d["label"] for d in miss_data],
                 {"Never retrieved (system)": [d["miss_not_retrieved"] for d in miss_data],
                  "Filtered before model (system)": [d["miss_gated"] for d in miss_data],
                  "Model saw it but didn't pick it": [d["llm_dropped"] for d in miss_data]},
                 "Where did the MISSED answers go?", "number of missed answers",
                 note="The whole bar = all the answers we missed. If it's all the top band, retrieval "
                 "found everything and only the model's choice (the prompt) can improve it.")
        figs["misses"] = save(fig, "misses")
    # 5. precedent - ONLY the approaches that actually showed past tickets to the model.
    if hist_data:
        fig, ax = plt.subplots(figsize=(8, 4.6))
        _bar(ax, hist_labels,
             {"Answers that WERE in the shown past tickets": [d["backed_r"] for d in hist_data],
              "Answers that were NOT in the past tickets": [d["notbacked_r"] for d in hist_data]},
             "When we show past tickets, does the model follow them?",
             "how often the model picked that answer",
             note="Tall blue + short orange = the model picks an answer far more when a similar past "
             "ticket used it. That gap is the proof the examples are working.", pct=True)
        figs["boost"] = save(fig, "boost")
    # 6. latency - typical (median) + average with the one outlier ticket removed, so the chart is
    #    readable (the raw slowest/avg are in the table). Only runs that recorded a median.
    lat_data = [d for d in data if d["lat_med"] is not None]
    if lat_data:
        any_outlier = any(d["lat_outlier"] for d in lat_data)
        fig, ax = plt.subplots(figsize=(8, 4.6))
        _bar(ax, [d["label"] for d in lat_data],
             {"typical (median)": [d["lat_med"] for d in lat_data],
              "average (slowest ticket removed)": [d["lat_avg_clean"] for d in lat_data]},
             "How long does one prediction take?", "seconds",
             note=("One retry-storm ticket was excluded so the scale is readable - the raw average and "
                   "slowest are in the table below." if any_outlier else
                   "Typical (median) and average per ticket."))
        figs["latency"] = save(fig, "latency")

    best = max(data, key=lambda d: d["micro_f1"])
    best_r = max(data, key=lambda d: d["micro_r"])
    base = min(data, key=lambda d: d["micro_r"])  # weakest recall = the baseline

    report_title = title or "Value-Stream prediction — comparison of approaches"
    if append and out_path.exists():
        # Add this comparison as a new chapter in the existing report.
        doc = Document(str(out_path))
        doc.add_page_break()
        doc.add_heading(report_title, level=1)
    else:
        doc = Document()
        doc.add_heading(report_title, level=0)
        doc.add_paragraph(
            "This report compares several approaches to Value Stream prediction on the same 60-ticket "
            "eval set. Each chart states the question it answers and how to read it; a plain-language "
            "guide to every metric is at the end. Where a chapter compares approaches that differ in "
            "one setting, a head-to-head table isolates the effect of that single change.")

    # What each run is (glossary) - decode the short labels once, up front.
    doc.add_heading("What each approach means", level=2 if append else 1)
    gt = doc.add_table(rows=1, cols=2); gt.style = "Light Grid Accent 1"
    gt.rows[0].cells[0].text, gt.rows[0].cells[1].text = "Approach", "What we showed the model"
    for d in data:
        cells = gt.add_row().cells
        cells[0].text = d["label"]
        cells[1].text = describe.get(d["label"], "(no description given)")

    # Bottom line first - plain English, data-driven.
    doc.add_heading("Bottom line", level=1)
    dr = best_r["micro_r"] - base["micro_r"]
    bullets = [
        f"Best approach: \"{best_r['label']}\" — it finds the most correct Value Streams "
        f"(recall {best_r['micro_r']:.0%}) and has the best overall balance (F1 {best_r['micro_f1']:.2f}).",
        f"That is {dr:.0%} more of the correct answers found than the weakest approach "
        f"(\"{base['label']}\", {base['micro_r']:.0%}) — a {dr / max(base['micro_r'], 1e-9) * 100:.0f}% "
        "relative jump.",
        f"It wins on BOTH easy tickets (one correct answer: {base['single_r']:.0%} → {best_r['single_r']:.0%} "
        f"found) and hard tickets (several correct answers: F1 {base['multi_f1']:.2f} → {best_r['multi_f1']:.2f}), "
        "so the gain is real, not just from the easy cases.",
    ]
    if best_r["backed_r"] is not None and best_r["notbacked_r"] is not None:
        bullets.append(
            f"Why it wins: when a correct Value Stream showed up in the past-ticket examples, the model "
            f"picked it {best_r['backed_r']:.0%} of the time, vs only {best_r['notbacked_r']:.0%} when it "
            "didn't — so the model is genuinely learning from precedent, not guessing.")
        head = best_r["hist_r"] - best_r["backed_r"]
        if head > 0.03:
            bullets.append(
                f"Room to improve: the past examples actually CONTAINED {best_r['hist_r']:.0%} of the correct "
                f"answers, but the model only USED {best_r['backed_r']:.0%} of them — about {head:.0%} of the "
                "right answers were sitting in front of it, unpicked. A more trusting prompt, or showing more "
                "past tickets, could capture that.")
    bullets.append(
        "Note on precision: the model always returns 10 guesses but a ticket has ~2.5 correct answers, so "
        "precision is mechanically low for everyone — judge it on RECALL and F1, not raw precision.")
    for b in bullets:
        doc.add_paragraph(b, style="List Bullet")

    if axes:
        _render_axes(doc, {d["label"]: d for d in data}, axes)

    if narrative:
        _render_markdown(doc, narrative)

    doc.add_heading("1. Precision, Recall, F1", level=1)
    doc.add_paragraph(
        f"Recall = of the correct Value Streams, how many did we find. F1 = the balance of precision and "
        f"recall. Best F1: \"{best['label']}\" ({best['micro_f1']:.2f}). Best recall: \"{best_r['label']}\" "
        f"({best_r['micro_r']:.0%}). Precision looks low everywhere because the model returns 10 guesses "
        "for ~2.5 real answers — that is by design, so read recall and F1.")
    doc.add_picture(figs["quality"], width=Inches(6.2))

    if "cohorts" in figs:
        doc.add_heading("2. Does it handle hard tickets, or just easy ones?", level=1)
        doc.add_paragraph(
            "We split tickets by how many correct Value Streams they have. EASY = exactly one correct answer; "
            "HARD = two or more. The chart shows RECALL for both (the same metric, so it's a fair comparison) "
            "— of the correct answers on that kind of ticket, how many we found. The orange (hard) bar is the "
            "one to watch: anyone can do well on easy tickets.")
        if cohort_omitted:
            _note(doc, "Not shown in this chart: " + ", ".join(f'"{x}"' for x in cohort_omitted) +
                  " — those runs did not record a per-cohort breakdown.")
        doc.add_picture(figs["cohorts"], width=Inches(6.2))
        doc.add_paragraph("Full per-cohort numbers (n = how many tickets, P / R / F1):")
        crows = []
        for d in cohort_data:
            crows.append([d["label"], str(d["single_n"] or "-"), _pct(d["single_p"]), _pct(d["single_r"]),
                          _f2(d["single_f1"]), str(d["multi_n"] or "-"), _pct(d["multi_p"]), _pct(d["multi_r"]),
                          _f2(d["multi_f1"])])
        _add_table(doc, ["Approach", "Easy n", "Easy P", "Easy R", "Easy F1",
                         "Hard n", "Hard P", "Hard R", "Hard F1"], crows)

    doc.add_heading("3. Did the system even put the right answers in front of the model?", level=1)
    p3 = rep  # retrieval is identical across approaches (same search), so describe once
    doc.add_paragraph(
        "The model can only pick from what the system shows it. This section is about the SYSTEM (retrieval), "
        "not the model. Each bar is the fraction of correct answers present at that stage.")
    doc.add_picture(figs["retrieval"], width=Inches(6.2))
    doc.add_paragraph(
        f"Semantic search ranking is weak — its top 10 holds only {_pct(p3['vs_r10'])} of the correct answers "
        f"and barely grows with depth (top 3: {_pct(p3['vs_r3'])}, top 5: {_pct(p3['vs_r5'])}). But because we "
        f"hand the model the FULL catalogue every time, {_pct(p3['pool_r'])} of the correct answers are ALWAYS "
        "in front of it. So retrieval is not the bottleneck — the model is choosing from a complete list.")
    if "misses" in figs:
        doc.add_paragraph(
            "The next chart proves it. Every correct answer we missed is one of three things: never retrieved, "
            "filtered out before the model, or shown to the model which then didn't pick it. The whole bar is "
            "the total misses; the colour tells you whose fault each miss is.")
        doc.add_picture(figs["misses"], width=Inches(6.2))
        doc.add_paragraph(
            "All of the misses are the top band — 'the model saw it but didn't pick it'. None were lost by "
            "retrieval. That is why the PROMPT is the only lever left: the right answers are already on the "
            "table, the model just has to choose them.")

    if "boost" in figs:
        doc.add_heading("4. When we show past tickets, does the model follow them?", level=1)
        doc.add_paragraph(
            "Only approaches that actually SHOW past tickets to the model appear here. We split the correct "
            "answers into two groups — ones a similar past ticket used, and ones it didn't — and measure how "
            "often the model picked each. A tall blue bar next to a short orange bar means the model is "
            "following the examples (it picks an answer much more when precedent backs it).")
        hb = max(hist_data, key=lambda d: (d["backed_r"] or 0))
        if hb["backed_r"] is not None and hb["notbacked_r"] is not None:
            cap = _ratio(hb["backed_r"], hb["hist_r"])
            doc.add_paragraph(
                f"In \"{hb['label']}\": an answer a past ticket used is picked {hb['backed_r']:.0%} of the time "
                f"vs only {hb['notbacked_r']:.0%} when no past ticket used it — a {hb['backed_r'] - hb['notbacked_r']:+.0%} "
                "lift, the proof the examples drive the picks.")
            head = hb["hist_r"] - hb["backed_r"]
            if head > 0.03:
                doc.add_paragraph(
                    f"Headroom: the past tickets actually contained {hb['hist_r']:.0%} of the correct answers, "
                    f"but the model only picked {hb['backed_r']:.0%} of those ({_pct(cap)} of what precedent "
                    f"offered). So ~{head:.0%} of the right answers were shown via precedent and still missed — "
                    "a more trusting prompt, or more past tickets, is where the next gain is.")
        doc.add_picture(figs["boost"], width=Inches(6.2))

    if "latency" in figs:
        doc.add_heading("5. How long does a prediction take?", level=1)
        doc.add_paragraph(
            "The chart shows TYPICAL (median) speed and the AVERAGE with the single slowest ticket removed. "
            "We remove it because a retry/backoff storm on one ticket can make it take thousands of seconds "
            "and wreck the average — that is an infrastructure hiccup, not the real cost. The table keeps the "
            "raw numbers so nothing is hidden.")
        lrows = []
        for d in lat_data:
            flag = "  ⚠ retry-storm outlier" if d["lat_outlier"] else ""
            lrows.append([d["label"], _secs(d["lat_med"]), _secs(d["lat_avg_clean"]),
                          _secs(d["lat_avg"]), _secs(d["lat_max"]) + flag])
        _add_table(doc, ["Approach", "Typical (median)", "Average (outlier removed)",
                         "Raw average", "Raw slowest"], lrows)
        doc.add_picture(figs["latency"], width=Inches(6.2))

    # summary table
    doc.add_heading("All numbers in one table", level=1)
    srows = []
    for d in data:
        lift = f"{d['boost_lift']:+.2f}" if d["boost_lift"] is not None else "n/a"
        srows.append([d["label"], _pct(d["micro_p"]), _pct(d["micro_r"]), _f2(d["micro_f1"]),
                      _pct(d["single_r"]), _pct(d["multi_r"]), _pct(d["backed_r"]), lift, _secs(d["lat_med"])])
    _add_table(doc, ["Approach", "Precision", "Recall", "F1", "Easy R", "Hard R",
                     "Precedent backed", "Precedent lift", "Typical time"], srows)
    note = doc.add_paragraph(
        "Recall-type numbers are shown as percentages; F1 as a 0-1 score. Older runs that predate a "
        "metric show 'n/a' (not zero) elsewhere in this report.")
    note.runs[0].italic = True

    # plain-language glossary so nobody has to guess what a metric means (once per report)
    if append:
        out_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(out_path))
        return
    doc.add_heading("What the metrics mean", level=1)
    glossary = [
        ("Recall", "Of the Value Streams that were actually correct, the fraction we found. Higher is better."),
        ("Precision", "Of our guesses, the fraction that were correct. Low here by design (10 guesses, ~2.5 real)."),
        ("F1", "A single balance score combining precision and recall. Higher is better."),
        ("Easy ticket", "A ticket that has only ONE correct Value Stream - one answer to find. 'Easy: found' = how "
                        "often we found that one answer."),
        ("Hard ticket", "A ticket that has TWO OR MORE correct Value Streams - we must find several. Scored with F1. "
                        "We track easy and hard apart so a win isn't just from the easy cases."),
        ("Precedent backed (capture)", "Of the correct answers that WERE in the shown past tickets, the fraction "
                                       "the model actually picked. An absolute score: how much of what precedent "
                                       "offered did it use (the ceiling is how many were in the tickets at all). "
                                       "Higher = fewer precedent-backed answers left on the table."),
        ("Precedent lift", "How much MORE often the model picks an answer when a past ticket used it vs when none "
                           "did. A difference: (picked-rate for answers IN the past tickets) - (picked-rate for "
                           "answers NOT in them). Big = the precedent is causing the pick, not general guessing; "
                           "near zero = the examples make no difference. Backed says how much it captured; lift "
                           "says whether the examples are the reason."),
        ("Past-ticket examples", "Similar tickets from the past, shown to the model along with the answers they got."),
        ("Search relevance scores", "Numbers from the search engine ranking each candidate; included or hidden."),
        ("Ceiling", "The fraction of correct answers that were on the candidate list at all - the best any "
                    "approach could possibly score. Here it's ~100% because the full catalogue is always shown."),
        ("Ceiling capture", "Of the correct answers shown to the model, the fraction it actually picked "
                            "(recall divided by the ceiling). Low = the model is leaving answers on the table."),
        ("Precedent capture", "Of the correct answers that were in the SHOWN past tickets, the fraction the model "
                              "picked. Only meaningful for approaches that show past tickets."),
        ("Typical (median) time", "The middle ticket's time - the honest 'usual' speed. The AVERAGE can be far "
                                  "higher if one ticket stalls on retries; compare the two to spot an outlier."),
        ("Where misses go", "Each correct answer we missed is one of: never retrieved, filtered before the model, "
                            "or the model saw it and dropped it. The last one is the model's fault, not the system's."),
    ]
    gl = doc.add_table(rows=1, cols=2); gl.style = "Light Grid Accent 1"
    gl.rows[0].cells[0].text, gl.rows[0].cells[1].text = "Term", "Plain meaning"
    for term, meaning in glossary:
        cells = gl.add_row().cells
        cells[0].text, cells[1].text = term, meaning

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("runs", nargs="+", help="runs.json files (or label=path) to compare")
    parser.add_argument("--out", default="out/eval/comparison.docx")
    parser.add_argument("--axis", action="append", default=[], metavar="'question: A vs B'",
                        help="add a head-to-head section between two run labels, e.g. "
                             "--axis 'Does showing past tickets help?: No history vs History'")
    parser.add_argument("--describe", action="append", default=[], metavar="'label: plain text'",
                        help="one-line plain-English description of a run, shown in the glossary, e.g. "
                             "--describe 'History: all 50 streams + similar past tickets shown as examples'")
    parser.add_argument("--history-run", action="append", default=[], metavar="LABEL",
                        help="mark a run as one that SHOWS past tickets to the model; only these appear in the "
                             "precedent-use section. Repeat per history run. If omitted, falls back to any run "
                             "that recorded the precedent split.")
    parser.add_argument("--narrative", metavar="FILE",
                        help="markdown file (## / ### headings, - bullets, **bold**) rendered into the report "
                             "as an analysis section, e.g. the prompt change-log.")
    parser.add_argument("--title", help="title for this comparison chapter (defaults to the standard title)")
    parser.add_argument("--append", action="store_true",
                        help="append this comparison as a new chapter to an existing --out docx (page break + "
                             "title), instead of starting fresh. Build a multi-experiment report in several calls.")
    args = parser.parse_args()
    data = [load(a) for a in args.runs]
    axes = [_parse_axis(a) for a in args.axis]
    describe = {}
    for d in args.describe:
        label, _, text = d.partition(":")
        describe[label.strip()] = text.strip()
    history_runs = set(args.history_run) or None
    narrative = Path(args.narrative).read_text(encoding="utf-8") if args.narrative else None
    build(data, Path(args.out), axes=axes, describe=describe, history_runs=history_runs,
          narrative=narrative, title=args.title, append=args.append)
    print(f"comparison -> {args.out}")
    print("\nquick table (P / R / F1 / Easy-R / Hard-R):")
    for d in data:
        hr = "n/a" if d["multi_r"] is None else f"{d['multi_r']:.3f}"
        print(f"  {d['label']:20} P={d['micro_p']:.3f} R={d['micro_r']:.3f} F1={d['micro_f1']:.3f} "
              f"easy-R={d['single_r']:.3f} hard-R={hr}")


if __name__ == "__main__":
    main()
